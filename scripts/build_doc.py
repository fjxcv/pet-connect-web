# -*- coding: utf-8 -*-
"""生成功能测试用例文档 —— 按模板格式：部分字段同行（4列），每用例一个独立表格"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LQ, RQ = '【', '】'
def Q(s): return f'{LQ}{s}{RQ}'

# ====== 辅助函数 ======
def set_shading(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    tcPr.append(s)

def set_valign(cell, a='top'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign'); v.set(qn('w:val'), a); tcPr.append(v)

def cell_text(cell, text, bold=False, fs=9, fn='宋体', align=None):
    for p in cell.paragraphs: p.clear()
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(str(text) if text else '')
    r.font.size = Pt(fs); r.font.name = fn; r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    r.bold = bold
    return p

# 表格结构：single=单字段整行(2列), double=双字段同行(4列)
TABLE_STRUCTURE = [
    {'type': 'single', 'label': '模块名',         'key': 'module',    'bold': True,  'w': 2.2},
    {'type': 'double', 'label1': '用例作者', 'key1': 'author',  'label2': 'BUG编号', 'key2': 'bug',    'w1': 1.6, 'w2': 1.4},
    {'type': 'double', 'label1': '测试员',   'key1': 'tester',  'label2': '测试日期', 'key2': 'date',   'w1': 1.6, 'w2': 1.4},
    {'type': 'double', 'label1': '测试类型', 'key1': 'type',    'label2': '测试工具', 'key2': 'tool',   'w1': 1.6, 'w2': 1.4},
    {'type': 'single', 'label': '用例ID',         'key': 'id',        'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '用例描述',       'key': 'desc',      'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '前驱条件',       'key': 'pre',       'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '测试步骤',       'key': 'steps',     'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '期待结果',       'key': 'expect',    'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '实际结果',       'key': 'actual',    'bold': False, 'w': 2.2},
    {'type': 'single', 'label': '备注',           'key': 'notes',     'bold': False, 'w': 2.2},
]

def add_one_case_table(doc, case, sub_module):
    """按照模板格式生成单个用例表格：同行字段合并为4列"""
    nrows = len(TABLE_STRUCTURE)
    table = doc.add_table(rows=nrows, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    label_shading = 'E8F5E9'

    for ri, spec in enumerate(TABLE_STRUCTURE):
        if spec['type'] == 'single':
            c0 = table.cell(ri, 0)
            c0.width = Cm(2.0)
            cell_text(c0, spec['label'], bold=spec.get('bold', False), fs=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_shading(c0, label_shading)
            c1 = table.cell(ri, 1)
            c1_merge = table.cell(ri, 3)
            c1.merge(c1_merge)
            c1.width = Cm(13.5)
            val = case.get(spec['key'], '')
            if spec['key'] == 'module':
                cell_text(c1, f'{case.get("module","")}——{sub_module}', fs=8)
            else:
                cell_text(c1, str(val) if val else '', fs=8)
            set_valign(c1, 'top')
        else:
            c0 = table.cell(ri, 0); c0.width = Cm(1.8)
            cell_text(c0, spec['label1'], bold=True, fs=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_shading(c0, label_shading)
            c1 = table.cell(ri, 1); c1.width = Cm(4.5)
            cell_text(c1, str(case.get(spec['key1'], '')), fs=8); set_valign(c1, 'top')
            c2 = table.cell(ri, 2); c2.width = Cm(1.8)
            cell_text(c2, spec['label2'], bold=True, fs=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_shading(c2, label_shading)
            c3 = table.cell(ri, 3); c3.width = Cm(5.4)
            cell_text(c3, str(case.get(spec['key2'], '')), fs=8); set_valign(c3, 'top')
    doc.add_paragraph()

# ====== 公共字段值 ======
M = '流浪宠物综合救助管理平台'  # module
A = '曹心如'                    # author/tester
D = '2026/6/24'                 # date
T = '前台功能测试'               # type
TL = '无'                       # tool

# ======================================================================
# 宠物领养模块 (18条，含完整备注)
# ======================================================================
ADOPT = [
    # (1) 列表展示
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-001','desc':'验证宠物列表页面默认加载和展示功能',
     'pre':'后端服务正常运行，数据库中存在可领养宠物数据（adoption_status=available, is_public=True）',
     'steps':'1. 访问 /pets，等待页面加载完成\n2. 观察宠物卡片列表的展示情况\n3. 检查每张卡片包含的信息项\n4. 查看页面顶部搜索框和筛选控件\n5. 查看页面统计数据',
     'expect':'1. 页面正常加载，以卡片网格展示所有可领养宠物\n2. 每张卡片含：照片缩略图、名称、品种、月龄、性别、地区、体型、健康状况\n3. 页面顶部有搜索框、物种/性别/年龄筛选下拉框\n4. 显示'+Q('共找到 N 只可领养宠物')+'统计\n5. 非登录用户也可正常访问',
     'actual':'','notes':'API: GET /pets/；AllowAny公开访问；列表仅展示is_public=True且adoption_status=available的宠物'},

    # (2) 搜索
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-002','desc':'验证宠物列表关键词搜索功能（按名称回车触发）',
     'pre':'数据库中有多条宠物数据（包含名称含'+Q('小白')+'的宠物）',
     'steps':'1. 在搜索框输入关键词'+Q('小白')+'，按回车\n2. 观察列表过滤结果\n3. 输入不存在的名称（如'+Q('xyz123')+'），再次搜索\n4. 清空搜索框按回车',
     'expect':'1. 输入后仅展示名称包含'+Q('小白')+'的宠物\n2. 无匹配时显示空状态提示'+Q('没有符合筛选条件的可领养宠物')+'\n3. 空状态提供'+Q('清除全部筛选条件')+'按钮\n4. 搜索通过回车触发（非实时）',
     'actual':'','notes':'API参数: search；回车确认搜索（非实时搜索）；空状态组件含清除筛选按钮'},

    # (3) 组合筛选
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-003','desc':'验证宠物列表多条件组合筛选功能',
     'pre':'数据库中存在多种物种/性别/年龄的宠物',
     'steps':'1. 选择物种='+Q('狗')+'，观察列表\n2. 再选择性别='+Q('母')+'\n3. 选择年龄='+Q('幼龄（0-6个月）')+'\n4. 结合关键词搜索\n5. 点击重置按钮',
     'expect':'1. 多条件同时生效（AND关系），结果准确\n2. 下拉选择立即触发筛选\n3. 统计数据随筛选实时更新\n4. 重置后所有条件清空，恢复全量',
     'actual':'','notes':'API参数: species/gender/age_min/age_max/search；多条件为AND关系；重置调用一键清除'},

    # (4) 附近搜索
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-004','desc':'验证附近搜索功能（地理位置 + Haversine距离计算）',
     'pre':'1. 浏览器允许地理位置权限\n2. 数据库中有带经纬度的宠物\n3. 高德逆地理编码API可用',
     'steps':'1. 点击'+Q('附近搜索')+'\n2. 浏览器弹窗点'+Q('允许')+'\n3. 选择半径（如5km），观察列表',
     'expect':'1. 获取GPS后高德逆编码得省/市\n2. Haversine公式计算距离，返回同省指定半径内宠物\n3. 卡片显示距离（如'+Q('距您3.2km')+'），按距离排序\n4. 支持切换半径（3/5/10/20/50km）',
     'actual':'','notes':'API: GET /pets/nearby/?lat=&lon=&radius_km=；Haversine公式；默认同省过滤same_province'},

    # (5) 详情展示
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-005','desc':'验证宠物详情页信息展示和Markdown描述渲染',
     'pre':'数据库中有完整宠物档案（含多张照片和Markdown格式description）',
     'steps':'1. 点击卡片进入 /pets/{id}\n2. 查看大图区域和'+Q('等待领养')+'标签\n3. 点击左右箭头翻看照片\n4. 查看详细信息区\n5. 滚动到描述区查看Markdown渲染',
     'expect':'1. 左侧大图叠加'+Q('等待领养')+'标签；右侧问候卡片含'+Q('带我回家')+'按钮\n2. 多图支持翻页，显示位置（如3/15），首末张箭头禁用\n3. 下方：种类/性别/年龄/体型/地区/健康状态\n4. Markdown标题/粗体/列表/链接正确渲染\n5. 面包屑：首页>领养列表>[宠物名]',
     'actual':'','notes':'API: GET /pets/{id}/；公开页面；前端Markdown渲染组件；翻页切换的是available列表中的前后宠物'},

    # (6) 不可领养提示
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-006','desc':'验证宠物详情页对不可领养宠物的提示展示',
     'pre':"某宠物 adoption_status='pending' 或 'adopted'（非available）",
     'steps':'1. 通过直接URL访问该宠物详情页 /pets/{id}\n2. 查看'+Q('带我回家')+'按钮区域\n3. 对比可领养宠物的详情页',
     'expect':'1. 宠物信息正常展示\n2. 不显示'+Q('带我回家')+'按钮，显示'+Q('该宠物当前不可领养')+'提示\n3. 照片上不是'+Q('等待领养')+'标签而是对应状态\n4. 用户无法发起领养',
     'actual':'','notes':'adoption_status非available时前端条件渲染：隐藏'+Q('带我回家')+'按钮，显示不可领养提示'},

    # (7) 带我回家-已登录
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-007','desc':'验证'+Q('带我回家')+'按钮——已登录跳转和未登录拦截',
     'pre':'1. 目标宠物adoption_status=available\n2. 准备已登录和未登录两种状态',
     'steps':'1. 未登录点击'+Q('带我回家')+'，观察跳转\n2. 登录后点击'+Q('带我回家')+'，观察跳转\n3. 查看申请页面内容',
     'expect':'1. 未登录->跳转/login，登录后回到/adopt/{petId}\n2. 已登录->直接跳转/adopt/{petId}\n3. 页面显示三步流程指示器：Step1问卷->Step2材料->Step3完成',
     'actual':'','notes':'前端ProtectedRoute组件保护/adopt/:petId路由；登录后redirect回原页面'},

    # (8) 问卷填写
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-008','desc':'验证领养申请Step1——问卷表单字段、条件逻辑和客户端校验',
     'pre':'用户已登录，位于/adopt/{petId}页面Step1',
     'steps':'1. 观察10个问卷字段展示\n2. 住宅类型选'+Q('其他')+'、户外区域选'+Q('其他')+'、养宠经验选'+Q('是')+'，检查条件框\n3. 留空必填点'+Q('下一步')+'，验证校验\n4. 电话填'+Q('abc')+'、年龄填200，验证格式\n5. 全部正确填写后点'+Q('下一步'),
     'expect':'1. 字段：姓名/地址/电话/微信/性别/年龄/住宅类型/住房情况/户外区域(多选)/同意/经验/变化\n2. 条件字段按逻辑显隐（其他->输入框，是->描述框）\n3. 必填提示；电话非法提示格式；年龄超范围提示0-150\n4. 校验通过后POST创建申请+POST提交问卷，进入Step2',
     'actual':'','notes':'API顺序: POST /adopt/applications/ -> POST /adopt/applications/{id}/questionnaire/；问卷答案存JSON'},

    # (9) 附件上传
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-009','desc':'验证领养申请Step2——附件上传、格式校验和提交',
     'pre':'已完成Step1，进入Step2材料上传页面',
     'steps':'1. 查看上传区域（身份证+居住证明）\n2. 选jpg图片上传身份证，选pdf上传居住证明\n3. 尝试上传.exe/.doc/超过5MB文件\n4. 缺材料点'+Q('提交')+'\n5. 全部上传后点'+Q('提交'),
     'expect':'1. 两段式上传：先uploadAPI.upload到/uploads/，再adoptAPI.addAttachment关联\n2. .exe/.doc拦截提示'+Q('仅支持jpg/png/pdf')+'\n3. >5MB拦截提示'+Q('文件大小不能超过5MB')+'\n4. 缺材料提示必传\n5. 上传成功显示文件名+预览，边框变绿',
     'actual':'','notes':'两段上传：uploadAPI.upload(file,subdir)->URL，adoptAPI.addAttachment关联；文件类型id_card/housing_proof'},

    # (10) 提交成功
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-010','desc':'验证领养申请Step3——提交成功页面和宠物状态变更',
     'pre':'已完成Step1问卷和Step2附件上传',
     'steps':'1. 进入Step3观察成功页面\n2. 点击'+Q('查看申请详情')+'\n3. 返回点'+Q('返回')+'\n4. 回到宠物列表查看该宠物状态',
     'expect':'1. 成功图标+'+Q('申请提交成功！')+'标题+'+Q('1-3个工作日')+'提示\n2. '+Q('查看申请详情')+'-> /my-applications/{id}\n3. '+Q('返回')+'-> /pets\n4. 宠物adoption_status已变更为pending',
     'actual':'','notes':'提交后宠物状态自动变更为pending（后端perform_create中处理）'},

    # (11) 重复申请
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-011','desc':'验证领养申请——重复提交和他人已申请的冲突拦截',
     'pre':'1. 用户已对某宠物提交pending申请\n2. 另有用户B已登录',
     'steps':'1. 用户再次提交同宠物申请\n2. 用户B尝试申请已被用户A申请的宠物',
     'expect':'1. 后端返回错误'+Q('您已提交过该宠物的领养申请')+'\n2. 用户B被拦截，提示'+Q('该宠物已有进行中的领养申请')+'\n3. 宠物状态保持pending不变',
     'actual':'','notes':'AdoptApplicationSerializer.validate检查：同用户+同宠物+pending/approved状态；宠物adoption_status必须available'},

    # (12) 我的申请列表
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-012','desc':'验证'+Q('我的申请与核验')+'列表展示和状态映射',
     'pre':'用户已登录，提交过多条不同状态的申请',
     'steps':'1. 点击'+Q('我的申请与核验')+'进入 /my-applications\n2. 逐一查看不同状态标签的颜色和文字\n3. 记录哪些状态显示'+Q('查看详情')+'按钮\n4. 空状态时查看提示',
     'expect':'1. 表格：宠物头像+名称+状态+时间+操作\n2. pending黄-待审核；rejected红-审核拒绝+详情按钮\n3. need_material黄-待补充\n4. approved+scheduled黄-待核验；approved+passed绿-核验通过\n5. approved+failed红-核验失败+详情按钮\n6. 仅拒绝和核验失败显示详情按钮',
     'actual':'','notes':'API: GET /adopt/applications/my/；前端getDisplayStatus映射7种状态-颜色对；详情按钮条件渲染'},

    # (13) 分页
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-013','desc':'验证'+Q('我的申请')+'列表分页功能',
     'pre':'用户申请记录超过10条',
     'steps':'1. 查看默认展示10条\n2. 点'+Q('下一页')+'加载第2页\n3. 点具体页码跳转\n4. 在首页检查'+Q('上一页')+'状态\n5. 在末页检查'+Q('下一页')+'状态',
     'expect':'1. 分页组件：页码列表+上/下页+总数\n2. 翻页数据正确加载\n3. 首页上页禁用；末页下页禁用\n4. 当前页码高亮',
     'actual':'','notes':'前端分页组件，每页10条；后端DRF PageNumberPagination'},

    # (14) 拒绝/失败详情
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-014','desc':'验证申请详情——查看审核拒绝和核验失败原因',
     'pre':'1. 有一条rejected申请（含audit_opinion）\n2. 有一条approved+verify_failed申请（含verify_note）',
     'steps':'1. 点rejected记录'+Q('查看详情')+'，查看页面\n2. 点failed记录'+Q('查看详情')+'，查看页面\n3. 对比两种详情页差异',
     'expect':'1. rejected：红色渐变头+'+Q('审核拒绝')+'+audit_opinion+audited_at\n2. failed：红色渐变头+'+Q('核验失败')+'+verify_note+verified_at\n3. 均通过getDetailType区分类型显示\n4. 返回按钮回/my-applications',
     'actual':'','notes':'getDetailType函数判断：核验失败优先于审核拒绝；仅申请人本人可查看（get_queryset按applicant过滤）'},

    # (15) 越权访问
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-015','desc':'验证申请详情——越权访问拦截',
     'pre':'用户A和用户B各有申请记录，用户A已登录',
     'steps':'1. 用户A通过URL尝试访问用户B的申请详情\n2. 观察系统反应',
     'expect':'1. 后端get_queryset按applicant过滤\n2. 返回403或'+Q('无权查看')+'\n3. 不泄露他人申请信息',
     'actual':'','notes':'后端retrieve/my时按applicant=request.user过滤；403通过DRF权限/查询集返回'},

    # (16) 未登录拦截
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-016','desc':'验证未登录用户无法访问需登录的领养页面',
     'pre':'用户未登录',
     'steps':'1. 直接访问 /my-applications\n2. 直接访问 /adopt/{petId}\n3. 观察系统反应',
     'expect':'1. 两个页面均被ProtectedRoute拦截\n2. 自动跳转 /login\n3. 不会短暂展示页面内容\n4. 登录后能回到目标页面',
     'actual':'','notes':'前端ProtectedRoute组件检查localStorage中的token；后端IsAuthenticated权限类'},

    # (17) 公开页面
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-017','desc':'验证宠物列表和详情为公开页面（未登录可正常访问）',
     'pre':'用户未登录',
     'steps':'1. 访问 /pets 列表页\n2. 点击卡片进入 /pets/{id} 详情页\n3. 查看是否显示管理员控件',
     'expect':'1. 列表和详情正常加载\n2. 所有公开宠物信息可见\n3. 看到'+Q('带我回家')+'按钮（但点击跳转登录）\n4. 不显示AdminManageBar管理员控件',
     'actual':'','notes':'后端views AllowAny权限；前端AdminManageBar通过role=admin条件渲染'},

    # (18) 权限汇总
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'ADOPT-018','desc':'验证领养模块三层权限体系——公开/登录/管理员',
     'pre':'准备未登录/普通用户/管理员三种身份',
     'steps':'1. 未登录：访问列表/详情/我的申请/申请页\n2. 普通用户：访问我的申请/提交申请/访问/add-pet\n3. 管理员：访问/add-pet编辑宠物',
     'expect':'1. 未登录：列表和详情可访问；我的申请和申请页跳转登录\n2. 普通用户：可提交申请和查看我的申请；/add-pet被AdminRoute拦截\n3. 管理员：可访问/add-pet进行宠物管理',
     'actual':'','notes':'三层权限：AllowAny(列表/详情) -> IsAuthenticated(申请) -> IsAdminRole(管理)；前端AdminRoute检查role'},
]

# ======================================================================
# 救助追踪模块 (18条，含完整备注)
# ======================================================================
RESCUE = [
    # (1) 列表展示
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-001','desc':'验证救助列表页面加载展示和描述文本生成',
     'pre':'数据库中有待救助（pending_rescue）案例',
     'steps':'1. 访问 /rescue\n2. 观察卡片内容\n3. 查看顶部按钮\n4. 对比不同属性案例的描述文本\n5. 刷新后确认已响应案例被过滤',
     'expect':'1. 卡片展示pending_rescue案例：照片/描述/编号/时间\n2. 描述含昵称+地点+体型+健康；怕人=是显示'+Q('胆子较小怕人')+'\n3. 顶部：我的救助/上报/查询三个按钮\n4. 已响应案例通过localStorage过滤\n5. 未登录也可浏览（AllowAny）',
     'actual':'','notes':'API: GET /rescue/cases/；列表AllowAny；前端buildDescription构造描述；localStorage存已响应ID列表'},

    # (2) 上报表单
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-002','desc':'验证上报救助——表单完整字段展示和必填校验',
     'pre':'用户已登录，进入 /rescue/report',
     'steps':'1. 观察表单字段和必填标记\n2. 空表单点提交，验证依次提示\n3. 联系方式填<5字符验证\n4. 不选体型/健康/怕人验证\n5. 全部正确填写后提交',
     'expect':'1. 必填：昵称(<=50字)/联系(>=5字符)/地址/体型/健康/怕人；选填：照片\n2. 逐字段提示'+Q('请填写')+'\n3. 联系<5字符提示'+Q('至少5个字符')+'\n4. 有红色星号标记必填\n5. 全部正确提交成功',
     'actual':'','notes':'前端handleSubmit校验；后端serializer.validate校验nickname/contact/discover_address必填；contact>=5'},

    # (3) 定位
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-003','desc':'验证上报救助——地理位置定位和异常处理',
     'pre':'用户已登录，在 /rescue/report',
     'steps':'1. 点'+Q('定位')+'，允许权限，观察地址填充\n2. 关闭权限点定位，观察提示\n3. 模拟超时场景',
     'expect':'1. 获取GPS后高德逆编码填充'+Q('省+市+区+街道')+'格式地址\n2. 经纬度存discover_lat/lng（6位小数）\n3. 拒绝权限提示'+Q('请手动输入')+'；超时提示'+Q('定位超时')+'(15s)\n4. 失败后仍可手动输入地址',
     'actual':'','notes':'浏览器Geolocation API + 高德逆地理编码(Key:0cd985d74a8143ce3e159294a37635a2)；超时15s；CoordinateField保留6位小数'},

    # (4) 照片上传
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-004','desc':'验证上报救助——照片上传和格式/大小校验',
     'pre':'用户已登录，在 /rescue/report',
     'steps':'1. 选jpg图片上传，观察缩略图\n2. 再选png上传\n3. 尝试上传>5MB图片\n4. 尝试上传.doc文件\n5. 点删除按钮移除照片\n6. 不上传照片直接提交',
     'expect':'1. 支持jpg/png/pdf，单文件<=5MB\n2. 上传成功显示缩略图预览\n3. >5MB拦截提示'+Q('文件大小不能超过5MB')+'\n4. .doc拦截提示'+Q('仅支持jpg/png/pdf')+'\n5. 删除按钮可移除\n6. 照片非必填，不上传可提交',
     'actual':'','notes':'照片前端校验格式+大小；photo_urls存JSON数组；可多张上传；非必填项'},

    # (5) 提交成功
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-005','desc':'验证上报救助——完整提交流程、编号生成和成功页',
     'pre':'用户已登录，已填写完毕',
     'steps':'1. 点'+Q('提交')+'，观察请求\n2. 查看成功页展示\n3. 记录救助编号\n4. 点'+Q('继续上报')+'\n5. 点'+Q('返回列表'),
     'expect':'1. POST /rescue/cases/，自动生成编号RES{YYYYMMDD}{3位流水号}\n2. 自动创建状态日志from=None->pending_rescue(remark='+Q('被发现上报')+')\n3. 成功页显示编号+提示保存\n4. 继续上报->空白表单，返回列表->/rescue',
     'actual':'','notes':'generate_rescue_no()：查询当天最大编号+1；create时perform_create自动创建首条RescueStatusLog'},

    # (6) 编号查询
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-006','desc':'验证救助编号精确查询和搜索历史记录',
     'pre':'数据库中有已知救助编号',
     'steps':'1. 点'+Q('查询')+'进入 /rescue/search\n2. 输入编号查询，点结果跳详情\n3. 输入不存在编号\n4. 多次搜索查看历史\n5. 点历史标签',
     'expect':'1. 搜索显示编号+状态，点结果跳/rescue/{id}\n2. 不存在显示'+Q('未找到')+'\n3. 空输入禁用查询按钮\n4. 历史最多3条存localStorage按用户隔离\n5. 点历史直接搜索；不重复保存',
     'actual':'','notes':'API: GET /rescue/cases/?rescue_no=XXX；iexact大小写不敏感；历史存localStorage[key=pet_connect_search_{userId}]'},

    # (7) 响应救助
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-007','desc':'验证响应救助——状态自动推进和联系方式展示',
     'pre':'1. 有一条pending_rescue案例\n2. 用户已登录未响应该案例',
     'steps':'1. 点'+Q('救助')+'按钮\n2. 观察弹窗\n3. 刷新列表确认案例消失\n4. 查看我的救助确认出现',
     'expect':'1. POST /rescue/cases/{id}/help/\n2. 用户加入helpers，help_date记录\n3. pending_rescue自动->in_medical\n4. 创建status_log\n5. 弹窗展示上报人contact\n6. 案例从列表消失（localStorage）',
     'actual':'','notes':'API: POST help；pending_rescue只能通过help推进；helpers为ManyToMany；首次响应设置help_date；operation_log记录'},

    # (8) 防重复+多人
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-008','desc':'验证响应救助——防重复响应和多用户协作',
     'pre':'1. 用户A已响应某案例\n2. 用户B已登录未响应',
     'steps':'1. 用户A再次点'+Q('救助')+'\n2. 用户B搜索该案例点'+Q('救助')+'\n3. 查看案例helpers列表',
     'expect':'1. 用户A重复响应返回400'+Q('已响应过')+'\n2. 用户B成功响应，加入helpers列表\n3. 状态保持in_medical（不重复推进）\n4. help_date不重复设置',
     'actual':'','notes':'后端检查user in helpers；前端localStorage双重防重复；helpers为ManyToMany可追加多人'},

    # (9) 详情时间线
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-009','desc':'验证救助详情——状态流转路径和阶段时间线展示',
     'pre':'案例经历 待救助->医疗中->恢复中，有status_logs和stage_records',
     'steps':'1. 进入 /rescue/{id}\n2. 查看顶部状态流转路径\n3. 点不同状态节点查看时间线\n4. 区分两种记录类型\n5. 悬停节点看预计时长',
     'expect':'1. 横向流程图：待救助->医疗中->恢复中\n2. 已过节点高亮，当前加粗，未达灰显\n3. 时间线：黄色图标=状态变更，蓝色图标=阶段记录\n4. 切换节点同步切换时间线\n5. tooltip显示预计停留天数',
     'actual':'','notes':'并行GET /rescue/cases/{id}/+ /rescue/cases/{id}/stage-records/；前端合并status_logs+stage_records按时间戳分组'},

    # (10) 我的救助列表
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-010','desc':'验证'+Q('我的救助')+'列表、状态筛选和分页',
     'pre':'用户响应过多条不同状态的案例',
     'steps':'1. 访问 /my-rescues\n2. 查看表格和操作按钮\n3. 筛选'+Q('医疗中')+'\n4. 查看终态记录按钮\n5. 测试分页（>10条）',
     'expect':'1. 表格：照片/编号/状态徽章/日期\n2. 非终态显示'+Q('更新状态')+'和'+Q('填写记录')+'按钮\n3. 筛选正确；全部恢复全量\n4. 终态不显示操作按钮\n5. 每页10条分页正常',
     'actual':'','notes':'API: GET /rescue/cases/?helped=true&status=；筛选排除pending_rescue/abandoned；RescueCasePagination(10/页)'},

    # (11) 更新状态
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-011','desc':'验证更新救助状态——正常流转和备注必填校验',
     'pre':'用户是helper，案例in_medical',
     'steps':'1. 点'+Q('更新状态')+'进入 /my-rescues/{id}/update-status\n2. 点'+Q('切换状态')+'看新状态\n3. 不填备注点'+Q('确认更新')+'\n4. 填备注后提交',
     'expect':'1. 新状态自动为STATUS_FLOW下一步(recovering)\n2. 空备注前后端双重拦截\n3. 正常提交后in_medical->recovering\n4. 创建status_log+operation_log',
     'actual':'','notes':'API: PATCH /rescue/cases/{id}/status/；STATUS_FLOW: in_medical->recovering->awaiting_adoption->rescued；备注前后端双重校验必填'},

    # (12) 状态机规则
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-012','desc':'验证救助状态机规则——禁止跳跃、回退和终态变更',
     'pre':'用户是helper，案例in_medical和rescued各一条',
     'steps':'1. 尝试in_medical直接跳awaiting_adoption\n2. 尝试回退到pending_rescue\n3. 尝试更新rescued终态案例\n4. 尝试通过update_status手动推进pending_rescue',
     'expect':'1. 跳跃拦截-'+Q('只能从recovering变更为')+'\n2. 回退拦截-'+Q('不允许回退')+'\n3. 终态拦截-'+Q('当前状态不允许变更')+'\n4. pending_rescue只能通过help推进',
     'actual':'','notes':'STATUS_FLOW只允许单向邻步流转；rescued/abandoned为终态不可变更；pending_rescue仅可通过help自动推进'},

    # (13) 端到端流转
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-013','desc':'验证救助状态完整流转链路端到端测试',
     'pre':'用户是helper，新案例从pending_rescue开始',
     'steps':'1. help->pending_rescue->in_medical\n2. update->in_medical->recovering\n3. update->recovering->awaiting_adoption\n4. update->awaiting_adoption->rescued\n5. 每步填备注，查看status_logs',
     'expect':'1. 完整链路走通\n2. 每步含备注+status_log记录\n3. status_logs完整记录from->to历史\n4. rescued后不可再变更',
     'actual':'','notes':'全链路: pending_rescue(help)->in_medical(update)->recovering(update)->awaiting_adoption(update)->rescued；5条status_logs'},

    # (14) 填写阶段记录
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-014','desc':'验证填写阶段记录——提交、历史列表和内容校验',
     'pre':'用户是helper，案例非终态',
     'steps':'1. 点'+Q('填写记录')+'进入 /my-rescues/{id}/stage-record\n2. 输入内容提交\n3. 查看历史列表\n4. 空内容提交\n5. 多次提交',
     'expect':'1. 新记录出现在历史顶部，输入框清空\n2. 历史按时间倒序，显示操作人+时间+内容\n3. 空内容前端拦截\n4. 可多次提交，每条独立',
     'actual':'','notes':'API: POST /rescue/cases/{id}/stage-records/；GET同URL获取列表；content前后端双重非空校验'},

    # (15) 未登录拦截
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-015','desc':'验证未登录用户无法执行救助写操作',
     'pre':'用户未登录',
     'steps':'1. 直接访问 /rescue/report\n2. 在列表点'+Q('救助')+'\n3. 直接访问 /my-rescues',
     'expect':'1. 上报/我的救助页面跳转/login\n2. 点击救助跳转/login\n3. 后端IsAuthenticated拒绝未认证请求',
     'actual':'','notes':'前端ProtectedRoute保护需登录路由；后端IsAuthenticated+IsActiveUser保护写操作'},

    # (16) 非helper拒绝
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-016','desc':'验证非救助人不能更新状态和填写阶段记录',
     'pre':'普通用户C（非helper非admin），某案例helpers不含C',
     'steps':'1. 用户C尝试PATCH更新状态\n2. 用户C尝试POST填写阶段记录',
     'expect':'1. 两次操作均返回403\n2. 提示'+Q('只有救助人或管理员才能')+'\n3. 案例数据和记录不变',
     'actual':'','notes':'update_status和stage_records均检查is_helper or is_admin；is_admin判断: is_superuser/is_staff/profile.role=admin'},

    # (17) 公开页面
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-017','desc':'验证未登录可浏览救助列表和详情但不显示操作按钮',
     'pre':'用户未登录',
     'steps':'1. 访问 /rescue 列表\n2. 进入 /rescue/{id} 详情\n3. 查看是否有操作按钮\n4. 尝试访问/my-rescues/{id}/stage-record',
     'expect':'1. 列表和详情正常加载\n2. 状态路径和时间线可见\n3. 不显示救助/更新状态等按钮\n4. 阶段记录页受保护跳转登录',
     'actual':'','notes':'列表/详情AllowAny；操作按钮组件需isAuthenticated才渲染；写操作路由受ProtectedRoute保护'},

    # (18) 封禁用户
    {'module':M,'author':A,'bug':'','tester':A,'date':D,'type':T,'tool':TL,
     'id':'RESCUE-018','desc':'验证被封禁用户无法执行救助写操作',
     'pre':'用户已登录但profile.status=1（封禁）',
     'steps':'1. 尝试上报POST /rescue/cases/\n2. 尝试响应POST help\n3. 尝试更新状态PATCH status\n4. 尝试填写记录POST stage-records',
     'expect':'1. IsActiveUser拦截所有写操作\n2. 返回403或'+Q('账户已被禁用')+'\n3. 浏览列表和详情不受影响',
     'actual':'','notes':'IsActiveUser权限类检查profile.status==0；读操作(AllowAny)不受封禁影响'},
]

# ================================================================
# 主程序
# ================================================================
def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)  # A4
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = '宋体'; style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10)

    # 封面
    for _ in range(5): doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('流浪宠物综合救助管理平台'); r.font.size = Pt(26); r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r.bold = True

    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run('宠物领养模块 & 救助追踪模块\n功能测试用例文档'); r.font.size = Pt(20)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph(); doc.add_paragraph()
    for line in [f'测试模块：宠物领养模块 / 救助追踪模块', f'用例作者：曹心如', f'测试人员：曹心如', f'测试日期：2026/6/24', f'测试类型：前台功能测试', f'测试工具：无', f'用例总数：{len(ADOPT)+len(RESCUE)} 条（领养{len(ADOPT)}+救助{len(RESCUE)}）']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(12)
    doc.add_page_break()

    h1 = doc.add_heading(f'一、宠物领养模块 —— 前台功能测试用例（共 {len(ADOPT)} 条）', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h1.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_paragraph()
    for case in ADOPT: add_one_case_table(doc, case, '宠物领养模块')

    doc.add_page_break()

    h2 = doc.add_heading(f'二、救助追踪模块 —— 前台功能测试用例（共 {len(RESCUE)} 条）', level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h2.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_paragraph()
    for case in RESCUE: add_one_case_table(doc, case, '救助追踪模块')

    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, '宠物领养与救助追踪-功能测试用例.docx')
    doc.save(out_path)
    print(f'SUCCESS: {out_path}')
    print(f'宠物领养模块：{len(ADOPT)} 条用例')
    print(f'救助追踪模块：{len(RESCUE)} 条用例')
    print(f'合计：{len(ADOPT)+len(RESCUE)} 条用例')

if __name__ == '__main__':
    main()

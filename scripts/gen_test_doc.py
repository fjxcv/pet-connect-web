# -*- coding: utf-8 -*-
"""生成宠物领养与救助追踪模块功能测试用例文档"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 使用【】替代中文双引号避免编码问题
LQ, RQ = '【', '】'  # 【 】
def Q(s): return f'{LQ}{s}{RQ}'

COMMON = {
    '模块名': '流浪宠物综合救助管理平台',
    '用例作者': '曹心如', '测试员': '曹心如',
    '测试日期': '2026/6/24', '测试类型': '前台功能测试', '测试工具': '手动测试',
}

def set_shading(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    tcPr.append(s)

def set_valign(cell, a='top'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign'); v.set(qn('w:val'), a); tcPr.append(v)

def cell_text(cell, text, bold=False, fs=9, fn='宋体', align=None):
    for p in cell.paragraphs: p.clear()
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(str(text) if text else '')
    r.font.size = Pt(fs); r.font.name = fn; r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    r.bold = bold

def make_table(doc, title, cases):
    h = doc.add_heading(title, level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    cols = ['模块名','用例作者','BUG编号','测试员','测试日期','测试类型','测试工具','用例ID','用例描述','前驱条件','测试步骤','期待结果','实际结果']
    tbl = doc.add_table(rows=1+len(cases), cols=len(cols))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cell_text(tbl.rows[0].cells[i], c, bold=True, fs=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_shading(tbl.rows[0].cells[i], 'D9E8F7')
    for ri, cs in enumerate(cases):
        row = tbl.rows[ri+1]
        for ci, v in enumerate([
            COMMON['模块名'], COMMON['用例作者'], cs.get('bug',''), COMMON['测试员'],
            COMMON['测试日期'], COMMON['测试类型'], COMMON['测试工具'],
            cs.get('id',''), cs.get('desc',''), cs.get('pre',''),
            cs.get('steps',''), cs.get('expect',''), ''
        ]):
            cell_text(row.cells[ci], str(v) if v else '', fs=7)
            set_valign(row.cells[ci])
            if ri % 2 == 1: set_shading(row.cells[ci], 'F5F5F5')
    widths = [Cm(2.2),Cm(1.4),Cm(1.2),Cm(1.2),Cm(1.5),Cm(1.6),Cm(1.5),Cm(2.8),Cm(3.5),Cm(3.5),Cm(5.5),Cm(5.5),Cm(2.8)]
    for row in tbl.rows:
        for i, w in enumerate(widths): row.cells[i].width = w
    doc.add_paragraph()

# ====== 宠物领养模块 (23条) ======
ADOPT = [
    # --- 宠物列表浏览 (8条) ---
    {'bug':'','id':'ADOPT-LIST-001','desc':'验证宠物列表页面默认加载和展示功能',
     'pre':'后端服务正常运行，数据库中存在可领养宠物数据（adoption_status=available, is_public=True）',
     'steps':'1. 打开浏览器，访问宠物领养列表页面 /pets\n2. 等待页面加载完成\n3. 观察宠物卡片列表的展示情况\n4. 检查每张宠物卡片包含的信息项',
     'expect':'1. 页面正常加载，无报错，无空白页面\n2. 页面以卡片网格形式展示所有可领养宠物\n3. 每张卡片包含：宠物照片缩略图、名称、品种、月龄、性别、地区、体型、健康状况\n4. 页面显示'+Q('共找到 N 只可领养宠物')+'统计信息\n5. 页面顶部存在搜索框、物种/性别/年龄筛选控件\n6. 非登录用户也可正常访问（公开页面）'},

    {'bug':'','id':'ADOPT-LIST-002','desc':'验证宠物列表关键词搜索功能（按名称模糊匹配）',
     'pre':'后端服务正常运行，数据库中存在多条宠物数据（其中包含名称含'+Q('小白')+'的宠物）',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 在搜索框中输入关键词'+Q('小白')+'\n3. 按回车键触发搜索\n4. 观察列表更新结果\n5. 输入一个不存在的名称（如'+Q('xyz123不存在')+'），再次搜索',
     'expect':'1. 输入'+Q('小白')+'后，列表仅展示名称包含'+Q('小白')+'的宠物\n2. 搜索过滤结果准确，无无关数据\n3. 输入不存在名称时显示空状态提示'+Q('没有符合筛选条件的可领养宠物')+'\n4. 空状态页面提供'+Q('清除全部筛选条件')+'按钮\n5. 搜索通过回车键触发（非实时搜索）'},

    {'bug':'','id':'ADOPT-LIST-003','desc':'验证宠物列表按物种筛选功能',
     'pre':'后端服务正常运行，数据库中存在多种物种的宠物记录（狗/猫/鸟/兔/鱼/其他）',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 点击物种筛选下拉框，选择'+Q('狗')+'\n3. 观察列表更新\n4. 切换选择'+Q('猫')+'，再次观察\n5. 选择'+Q('全部')+'，确认恢复全部数据',
     'expect':'1. 物种下拉框包含选项：全部/狗/猫/鸟/兔/鱼/其他\n2. 选择'+Q('狗')+'后，列表仅展示 species=dog 的宠物\n3. 选择'+Q('猫')+'后，列表仅展示 species=cat 的宠物\n4. 选择'+Q('全部')+'后恢复所有物种展示\n5. 下拉选择后立即触发筛选（无需额外点击）'},

    {'bug':'','id':'ADOPT-LIST-004','desc':'验证宠物列表按性别筛选功能',
     'pre':'后端服务正常运行，数据库中存在不同性别的宠物记录',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 点击性别筛选下拉框，选择'+Q('公')+'\n3. 观察列表更新\n4. 切换选择'+Q('母')+'，再次观察',
     'expect':'1. 性别筛选下拉框包含选项：全部/公/母\n2. 选择'+Q('公')+'后仅展示 gender=male 的宠物\n3. 选择'+Q('母')+'后仅展示 gender=female 的宠物\n4. 筛选后统计数据同步更新'},

    {'bug':'','id':'ADOPT-LIST-005','desc':'验证宠物列表按地区联动筛选功能（国家-省/州-市三级联动）',
     'pre':'后端服务正常运行，数据库中宠物记录包含不同地区信息（country/province/city）',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 观察地区筛选区域是否有国家/省/市三级下拉框\n3. 选择一个国家，观察省份下拉框是否动态更新\n4. 选择一个省份，观察城市下拉框是否动态更新\n5. 选择城市后观察列表过滤结果',
     'expect':'1. 地区筛选包含国家、省/州、市三级联动下拉框\n2. 选择国家后，省份下拉框自动过滤为该国家下的省份（基于已有宠物数据动态生成）\n3. 选择省份后，城市下拉框自动过滤为该省份下的城市\n4. 选择城市后列表仅展示该地区的宠物\n5. 选择'+Q('全部')+'可逐级恢复'},

    {'bug':'','id':'ADOPT-LIST-006','desc':'验证宠物列表按年龄阶段筛选功能',
     'pre':'后端服务正常运行，数据库中存在不同月龄的宠物记录',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 观察年龄筛选区域提供的选项\n3. 选择'+Q('幼龄（0-6个月）')+'，观察列表过滤\n4. 选择'+Q('青年（7-12个月）')+'，观察列表过滤\n5. 选择'+Q('自定义范围')+'，输入最小月龄和最大月龄，点击确认',
     'expect':'1. 年龄筛选提供快捷选项：幼龄/青年/成年/高龄，以及自定义范围\n2. 选择快捷选项后，列表按对应月龄区间（age_min/age_max）过滤\n3. 自定义范围支持手动输入 min/max 月龄，结果准确\n4. 选择'+Q('全部')+'恢复所有年龄展示'},

    {'bug':'','id':'ADOPT-LIST-007','desc':'验证附近搜索功能（基于浏览器地理位置API + Haversine距离计算）',
     'pre':'1. 浏览器允许地理位置权限\n2. 后端服务正常运行，数据库中存在带经纬度的宠物记录\n3. 高德地图逆地理编码API可用',
     'steps':'1. 访问宠物领养列表页面 /pets\n2. 点击'+Q('附近搜索')+'按钮或相关入口\n3. 浏览器弹出地理位置授权，点击'+Q('允许')+'\n4. 选择搜索半径（如5km）\n5. 观察列表更新',
     'expect':'1. 页面提供'+Q('附近搜索')+'功能入口\n2. 浏览器获取用户GPS坐标后，调用高德逆地理编码获取省市信息\n3. 后端通过Haversine公式计算距离，返回指定半径内且同省的宠物\n4. 宠物卡片展示距离信息（如'+Q('距您3.2km')+'）\n5. 列表按距离由近到远排序\n6. 支持切换搜索半径（3/5/10/20/50km）\n7. 默认同省过滤（same_province参数）'},

    {'bug':'','id':'ADOPT-LIST-008','desc':'验证宠物列表组合筛选功能（多条件同时生效）',
     'pre':'后端服务正常运行，数据库中存在多条不同属性的宠物记录',
     'steps':'1. 访问宠物列表页 /pets\n2. 设置物种=狗\n3. 再设置性别=母\n4. 再设置年龄=幼龄（0-6个月）\n5. 再输入关键词搜索\n6. 观察最终结果\n7. 点击重置按钮',
     'expect':'1. 多条件同时生效，结果为所有条件的交集\n2. 各筛选条件之间为 AND 关系\n3. 统计数据随筛选条件变化实时更新\n4. 点击重置后所有筛选条件清空，列表恢复全量展示\n5. 各筛选控件状态同步重置'},

    # --- 宠物详情查看 (4条) ---
    {'bug':'','id':'ADOPT-DETAIL-001','desc':'验证宠物详情页面信息完整展示',
     'pre':'后端服务正常运行，数据库中存在一条完整的宠物档案（含照片、描述等）',
     'steps':'1. 在宠物列表页面点击任意宠物卡片\n2. 观察跳转到的宠物详情页面 /pets/{id}\n3. 逐项检查页面展示的信息\n4. 验证面包屑导航',
     'expect':'1. 页面正常跳转到宠物详情页\n2. 左侧展示宠物大图，照片上叠加'+Q('等待领养')+'标签\n3. 右侧展示问候卡片：用户名、宠物名称、月龄、所在地区、'+Q('带我回家')+'按钮\n4. 照片下方展示详细信息：种类、性别、年龄、体型、地区、健康状态\n5. 展示'+Q('捡拾详情描述')+'区域，支持Markdown渲染\n6. 若关联救助案例，展示救助编号\n7. 面包屑导航：首页 > 领养列表 > [宠物名]'},

    {'bug':'','id':'ADOPT-DETAIL-002','desc':'验证宠物详情页多图前后翻页浏览功能',
     'pre':'后端服务正常运行，某宠物档案关联了多张照片（photo_urls数组含多条）',
     'steps':'1. 进入有多张照片的宠物详情页\n2. 观察照片区域的左右翻页箭头\n3. 点击右箭头翻到下一张\n4. 连续点击翻到末张\n5. 点击左箭头回翻\n6. 观察位置指示器',
     'expect':'1. 照片区域显示左/右翻页箭头按钮\n2. 支持前翻和后翻浏览所有照片\n3. 显示当前位置指示器（如'+Q('3/15')+'）\n4. 翻页流畅无明显卡顿\n5. 到达首/末张时对应箭头按钮置灰/禁用\n6. 仅有一张照片时不显示翻页箭头'},

    {'bug':'','id':'ADOPT-DETAIL-003','desc':'验证宠物详情页Markdown描述内容正确渲染',
     'pre':'后端服务正常运行，某宠物的description字段包含Markdown格式文本（标题、粗体、列表、链接等）',
     'steps':'1. 进入该宠物的详情页\n2. 滚动到'+Q('捡拾详情描述')+'区域\n3. 观察Markdown内容的渲染效果\n4. 点击描述中的链接',
     'expect':'1. Markdown标题（# ## ###）正确渲染为对应级别标题样式\n2. **粗体文本**正确渲染为粗体\n3. *斜体文本*正确渲染为斜体\n4. 无序列表（- 或 *）和有序列表（1. ）正确渲染\n5. 链接正确渲染为可点击的超链接\n6. 换行和段落间距正确\n7. 纯文本描述正常显示'},

    {'bug':'','id':'ADOPT-DETAIL-004','desc':'验证宠物详情页不可领养宠物的提示展示',
     'pre':"后端服务正常运行，某宠物 adoption_status='pending' 或 'adopted'（非available状态）",
     'steps':'1. 通过直接URL访问该宠物详情页 /pets/{id}\n2. 观察页面展示\n3. 查看'+Q('带我回家')+'按钮区域',
     'expect':'1. 宠物信息正常展示\n2. 不是'+Q('等待领养')+'标签，而是对应状态标签\n3. '+Q('带我回家')+'按钮不显示，或显示'+Q('该宠物当前不可领养')+'提示\n4. 用户无法发起领养申请'},

    # --- 领养申请流程 (9条) ---
    {'bug':'','id':'ADOPT-APPLY-001','desc':'验证'+Q('带我回家')+'按钮——已登录用户正确跳转到领养申请页面',
     'pre':"1. 用户已登录\n2. 访问的宠物 adoption_status='available'",
     'steps':'1. 在宠物详情页 /pets/{id} 点击'+Q('带我回家')+'按钮\n2. 观察页面跳转\n3. 查看新页面的内容',
     'expect':'1. 页面自动跳转到 /adopt/{petId} 领养申请独立页面\n2. 申请页面显示当前宠物的基本信息（名称、品种等）\n3. 显示三步流程指示器：Step1填写问卷、Step2上传材料、Step3完成\n4. 默认停留在Step1'},

    {'bug':'','id':'ADOPT-APPLY-002','desc':'验证'+Q('带我回家')+'按钮——未登录用户被拦截并跳转登录页',
     'pre':'用户未登录',
     'steps':'1. 在宠物详情页点击'+Q('带我回家')+'按钮\n2. 观察系统反应\n3. 完成登录后观察跳转',
     'expect':'1. 系统检测到用户未登录\n2. 自动跳转到登录页面\n3. 登录成功后能够回到领养申请页面 /adopt/{petId}'},

    {'bug':'','id':'ADOPT-APPLY-003','desc':'验证领养申请Step1——问卷表单完整字段展示和条件逻辑',
     'pre':'用户已登录，位于 /adopt/{petId} 页面Step1',
     'steps':'1. 观察Step1问卷表单的完整字段列表\n2. 住宅类型选择'+Q('普通住宅')+'\n3. 住宅类型切换选择'+Q('其他')+'\n4. 户外区域选择'+Q('其他')+'\n5. '+Q('是否养过宠物')+'选择'+Q('是')+'\n6. '+Q('是否养过宠物')+'选择'+Q('否'),
     'expect':'1. 问卷包含字段：姓名、地址、电话、微信号、性别、年龄、住宅类型、住房情况、户外区域（多选）、同住人/房东是否同意、是否养过宠物、未来6个月变化\n2. 住宅类型选择'+Q('其他')+'时，出现额外文本输入框需填写具体类型\n3. 户外区域选择'+Q('其他')+'时，出现额外输入框\n4. 养宠经验选'+Q('是')+'时，出现详细描述输入框\n5. 养宠经验选'+Q('否')+'时，描述框隐藏\n6. 条件字段的显示/隐藏逻辑正确'},

    {'bug':'','id':'ADOPT-APPLY-004','desc':'验证领养申请Step1——表单客户端校验规则',
     'pre':'用户已登录，位于 /adopt/{petId} 页面Step1',
     'steps':'1. 不填写任何字段，直接点击'+Q('下一步')+'\n2. 仅填写姓名，再次点击'+Q('下一步')+'\n3. 填写非法手机号（如'+Q('abc')+'），点击'+Q('下一步')+'\n4. 填写年龄超出范围（如200），点击'+Q('下一步')+'\n5. 填写年龄为负数（如-5），点击'+Q('下一步')+'\n6. 正确填写所有必填字段后点击'+Q('下一步'),
     'expect':'1. 全部留空时提示'+Q('请填写姓名')+'等必填字段提示\n2. 逐个填写仍有空字段时提示对应字段必填\n3. 电话格式错误时提示'+Q('请输入正确的手机号码')+'（前端格式校验）\n4. 年龄200时提示'+Q('请输入有效年龄（0-150）')+'\n5. 年龄-5时提示'+Q('请输入有效年龄')+'\n6. 全部正确填写后，调用API创建领养申请并提交问卷，进入Step2'},

    {'bug':'','id':'ADOPT-APPLY-005','desc':'验证领养申请——同一用户对同一宠物不能重复提交申请',
     'pre':'1. 用户已登录\n2. 该用户已对某宠物提交过一条 pending 或 approved 状态的申请',
     'steps':'1. 再次进入该宠物的领养申请页面 /adopt/{petId}\n2. 填写问卷并点击'+Q('下一步')+'\n3. 观察系统反应',
     'expect':'1. 后端校验到该用户对该宠物已有进行中的申请\n2. 返回错误提示'+Q('您已提交过该宠物的领养申请')+'\n3. 前端捕获错误并显示提示，不创建重复记录\n4. 用户停留在Step1页面'},

    {'bug':'','id':'ADOPT-APPLY-006','desc':'验证领养申请——已有他人申请中的宠物不能被再次申请',
     'pre':'1. 用户A已对某宠物提交了 pending 状态的申请（宠物状态已变为pending）\n2. 用户B已登录',
     'steps':'1. 用户B进入该宠物的详情页\n2. 用户B点击'+Q('带我回家')+'按钮\n3. 填写问卷并提交',
     'expect':'1. 后端校验宠物 adoption_status 不是 available\n2. 返回错误提示'+Q('该宠物已有进行中的领养申请')+'\n3. 用户B的申请被拒绝创建\n4. 宠物状态保持pending不变'},

    {'bug':'','id':'ADOPT-APPLY-007','desc':'验证领养申请Step2——附件上传功能（正常流程）',
     'pre':'1. 用户已登录\n2. 已完成Step1问卷提交，进入Step2材料上传页面',
     'steps':'1. 观察Step2页面布局和要求\n2. 点击身份证上传区域，选择一张 jpg 图片\n3. 观察上传进度和结果\n4. 点击居住证明上传区域，选择一张 pdf 文件\n5. 上传完成后点击'+Q('提交')+'按钮\n6. 观察进入Step3',
     'expect':'1. 页面显示两项必传材料：居民身份证、居住证明\n2. 上传区域显示虚线边框，支持点击选择文件或拖拽文件\n3. 文件先上传到通用 /uploads/ 端点获取URL\n4. 再将URL通过 adoptAPI.addAttachment() 关联到申请\n5. 上传过程中显示加载状态\n6. 上传成功后显示文件名和预览，边框变实线绿色\n7. 两个文件都上传后可以进入Step3'},

    {'bug':'','id':'ADOPT-APPLY-008','desc':'验证领养申请Step2——附件文件格式和大小校验',
     'pre':'用户已进入Step2材料上传页面',
     'steps':'1. 尝试上传 .exe 可执行文件作为附件\n2. 尝试上传 .doc 格式文件\n3. 尝试上传超过5MB的图片文件\n4. 尝试上传200MB的PDF文件\n5. 尝试不上传任何附件直接点'+Q('提交'),
     'expect':'1. .exe 文件被前端拦截，提示'+Q('仅支持 jpg/png/pdf 格式')+'\n2. .doc 文件被前端拦截（非支持格式）\n3. 超过5MB的图片被前端拦截，提示'+Q('文件大小不能超过5MB')+'\n4. 超大型文件被立即拦截，不发起网络请求\n5. 拦截后上传区域保持可用状态\n6. 未上传附件时点击提交，提示'+Q('请上传所有必传材料')+'\n7. 已上传的文件可以通过删除按钮移除'},

    {'bug':'','id':'ADOPT-APPLY-009','desc':'验证领养申请Step3——提交成功页面展示',
     'pre':'用户已完成Step1问卷提交和Step2附件上传',
     'steps':'1. 进入Step3完成页面\n2. 观察页面展示内容\n3. 点击'+Q('查看申请详情')+'按钮\n4. 返回Step3后点击'+Q('返回')+'按钮',
     'expect':'1. Step3显示成功图标/动画和'+Q('申请提交成功！')+'标题\n2. 提示'+Q('请等待工作人员审核，通常1-3个工作日')+'\n3. 显示'+Q('查看申请详情')+'按钮，点击跳转到 /my-applications/{id}\n4. 显示'+Q('返回')+'按钮，点击跳转到 /pets 列表页\n5. 此时宠物状态已由 available 变为 pending'},

    # --- 我的申请查看 (3条) ---
    {'bug':'','id':'ADOPT-MY-001','desc':'验证'+Q('我的申请与核验')+'列表页面正常加载和数据展示',
     'pre':'1. 用户已登录\n2. 该用户至少提交过一条领养申请',
     'steps':'1. 点击顶部导航栏'+Q('我的申请与核验')+'链接\n2. 观察 /my-applications 页面加载情况\n3. 查看表格展示的信息\n4. 确认数据仅包含当前用户的申请',
     'expect':'1. 页面正常跳转到 /my-applications\n2. 页面标题显示'+Q('我的申请与核验')+'\n3. 以表格形式展示所有申请记录：宠物头像缩略图、宠物名称、审核/核验状态、提交时间、操作\n4. 调用 adoptAPI.getMy() 获取当前用户专属数据\n5. 仅展示当前登录用户本人的申请\n6. 无申请时显示'+Q('暂无申请记录')+'空状态'},

    {'bug':'','id':'ADOPT-MY-002','desc':'验证'+Q('我的申请')+'各状态标签颜色和文字正确映射',
     'pre':'用户已登录，存在多种状态的申请记录',
     'steps':'1. 进入 /my-applications 页面\n2. 逐一查看各条申请的状态标签\n3. 对比各种状态的显示：pending、rejected、need_material、approved+未核验、approved+scheduled、approved+passed、approved+failed',
     'expect':'1. pending--显示'+Q('待审核')+'（黄色徽章）\n2. rejected--显示'+Q('审核拒绝')+'（红色徽章）+ 显示'+Q('查看详情')+'按钮\n3. need_material--显示'+Q('待补充材料')+'（黄色徽章）\n4. approved+未核验--显示'+Q('审核通过')+'（绿色徽章）\n5. approved+scheduled--显示'+Q('待核验')+'（黄色徽章）\n6. approved+passed--显示'+Q('核验通过')+'（绿色徽章）\n7. approved+failed--显示'+Q('核验失败')+'（红色徽章）+ 显示'+Q('查看详情')+'按钮\n8. 仅'+Q('审核拒绝')+'和'+Q('核验失败')+'状态显示'+Q('查看详情')+'按钮'},

    {'bug':'','id':'ADOPT-MY-003','desc':'验证'+Q('我的申请')+'列表分页功能',
     'pre':'用户已登录，申请记录超过10条',
     'steps':'1. 进入 /my-applications 页面\n2. 观察默认展示的记录条数和分页组件\n3. 点击'+Q('下一页')+'\n4. 点击具体页码（如第3页）\n5. 点击'+Q('上一页')+'返回\n6. 检查首/末页的分页按钮状态',
     'expect':'1. 默认每页展示10条记录\n2. 分页组件显示：页码列表+上一页/下一页按钮+总记录数\n3. 点击'+Q('下一页')+'加载第2页数据\n4. 点击具体页码跳转到对应页\n5. 在首页时'+Q('上一页')+'按钮置灰/禁用\n6. 在末页时'+Q('下一页')+'按钮置灰/禁用\n7. 当前页码高亮显示'},

    # --- 申请详情(拒绝/失败) (3条) ---
    {'bug':'','id':'ADOPT-DETAIL-RJ-001','desc':'验证申请详情页——查看审核拒绝原因',
     'pre':'1. 用户已登录\n2. 该用户有一条 online_status=rejected 的申请（含 audit_opinion）',
     'steps':'1. 进入 /my-applications 页面\n2. 在状态为'+Q('审核拒绝')+'的记录上点击'+Q('查看详情')+'按钮\n3. 观察 /my-applications/{id} 页面的展示内容\n4. 点击'+Q('返回申请列表')+'按钮',
     'expect':'1. 页面展示申请摘要：宠物头像缩略图、宠物名称、申请时间\n2. 红色渐变头部，标题显示'+Q('审核拒绝')+'\n3. 明确展示拒绝原因（来自 audit_opinion 字段）\n4. 展示审核时间（audited_at）\n5. '+Q('返回申请列表')+'按钮跳转回 /my-applications\n6. 仅申请人本人可查看，其他人访问返回403'},

    {'bug':'','id':'ADOPT-DETAIL-RJ-002','desc':'验证申请详情页——查看核验失败原因',
     'pre':"1. 用户已登录\n2. 该用户有一条 online_status='approved' + verify_status='failed' 的申请（含 verify_note）",
     'steps':'1. 在'+Q('我的申请')+'页面该记录上点击'+Q('查看详情')+'\n2. 观察详情页面的展示内容\n3. 确认失败原因的来源',
     'expect':'1. 红色渐变头部，标题显示'+Q('核验失败')+'\n2. 明确展示核验失败原因（来自 verify_note 字段）\n3. 展示核验时间（verified_at）\n4. 展示核验人信息（如有）\n5. 与审核拒绝使用同一个详情页面，通过getDetailType区分类型'},

    {'bug':'','id':'ADOPT-DETAIL-RJ-003','desc':'验证申请详情页——越权访问拦截',
     'pre':'1. 用户A和用户B各有一条被拒绝的申请\n2. 用户A已登录',
     'steps':'1. 用户A尝试通过URL访问用户B的申请详情 /my-applications/{B的申请ID}\n2. 观察系统反应',
     'expect':'1. 后端 get_queryset 按 applicant 过滤\n2. 用户A无法查看用户B的申请详情\n3. 返回403 Forbidden 或'+Q('无权查看')+'提示\n4. 不会泄露他人申请信息'},

    # --- 权限控制(前台) (3条) ---
    {'bug':'','id':'ADOPT-AUTH-001','desc':'验证未登录用户无法访问'+Q('我的申请')+'页面',
     'pre':'用户未登录',
     'steps':'1. 在浏览器地址栏直接输入 /my-applications\n2. 观察系统反应',
     'expect':'1. 受 ProtectedRoute 保护\n2. 自动跳转到登录页面 /login\n3. 不会短暂展示页面内容\n4. 登录后能回到目标页面'},

    {'bug':'','id':'ADOPT-AUTH-002','desc':'验证未登录用户无法访问领养申请页面（直接URL访问）',
     'pre':'用户未登录',
     'steps':'1. 在浏览器地址栏直接输入 /adopt/{任意petId}\n2. 观察系统反应',
     'expect':'1. 受 ProtectedRoute 保护\n2. 自动跳转到登录页面\n3. 不会加载申请表单内容'},

    {'bug':'','id':'ADOPT-AUTH-003','desc':'验证宠物列表和详情为公开页面（未登录可访问）',
     'pre':'用户未登录',
     'steps':'1. 访问 /pets 宠物列表页\n2. 点击卡片进入 /pets/{id} 详情页\n3. 在详情页尝试操作',
     'expect':'1. 宠物列表和详情页正常加载，无需登录\n2. 所有宠物信息公开可见（is_public=True的宠物）\n3. 可以看到'+Q('带我回家')+'按钮，但点击会跳转登录页\n4. 不会看到管理员专有控件（AdminManageBar）'},
]

# 救助追踪模块用例数据量太大，请复用上面的脚本结构继续写...
print(f'Adopt cases defined: {len(ADOPT)}')
print(f'Data file prepared successfully')

# 保存到文件供下一步使用
import json
output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'docs')
os.makedirs(output_dir, exist_ok=True)

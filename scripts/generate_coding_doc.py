#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成宠物领养与救助追踪模块编码文档（Word格式）
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== 全局样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置默认段落间距
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_code_block(doc, code_text, file_path=None):
    """添加源代码块"""
    if file_path:
        p = doc.add_paragraph()
        run = p.add_run(f'📁 文件路径：{file_path}')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 100, 0)

    # 将代码以等宽字体添加
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(50, 50, 50)
        # 设置等宽字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 空行分隔
    doc.add_paragraph()


def add_subsection(doc, title, content):
    """添加子章节"""
    doc.add_heading(title, level=2)
    for paragraph in content.split('\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
        else:
            doc.add_paragraph()


# ===== 文档标题 =====
title = doc.add_heading('宠物领养与救助追踪模块编码文档', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目名称：暖爪救助 (PetConnect)')
doc.add_paragraph('技术栈：Django REST Framework + React')
doc.add_paragraph('文档生成日期：2026年6月23日')
doc.add_paragraph('')

# ===== 目录概述 =====
doc.add_heading('模块概述', level=1)
doc.add_paragraph(
    '本文档详细记录宠物领养申请与审核模块（5个子功能）和救助追踪模块（5个子功能）的实现细节，'
    '包括后端（Django REST API）和前端（React）的完整源代码、类结构、调用关系。'
)
doc.add_paragraph('后端代码位于 backend/pets/ 和 backend/rescue/ 目录下。')
doc.add_paragraph('前端代码位于 frontend/src/ 目录下，API 封装在 frontend/src/api/，页面在 frontend/src/pages/。')

# =============================================
# 模块一：宠物领养申请与审核模块
# =============================================
doc.add_heading('模块一：宠物领养申请与审核模块', level=1)

# ===== 1.1 宠物列表浏览 =====
doc.add_heading('1 宠物列表浏览功能', level=1)

doc.add_heading('1.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('文件：backend/pets/views.py 第19-71行 — PetProfileViewSet 类')
doc.add_paragraph('''
类名：PetProfileViewSet(viewsets.ModelViewSet)

属性：
- queryset = PetProfile.objects.all()  — 默认查询集
- serializer_class = PetProfileSerializer  — 关联的序列化器

方法：
- get_permissions(self)  — 第23-28行，动态权限控制：list/retrieve 公开访问，my_pets 需登录，其余需管理员
- get_queryset(self)  — 第30-62行，多条件筛选查询：物种/性别/体型/领养状态/搜索/地区/年龄范围/公开状态
- my_pets(self, request)  — 第64-71行，自定义 action，获取当前用户发布的领养宠物

URL 路由（backend/pets/urls.py 第5-8行）：通过 DefaultRouter 注册 PetProfileViewSet，由框架自动生成 RESTful 端点。''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('文件：frontend/src/pages/PetList.js 第165-754行 — PetList 组件')
doc.add_paragraph('''
组件名：PetList (React 函数组件)

核心 Hook：
- useState: pets(列表), loading, error, search/speciesFilter/genderFilter/ageRangeFilter/locationFilter(筛选条件)
- useEffect: 根据筛选条件变化自动请求宠物列表
- useCallback: buildApiParams() 构建 API 请求参数

主要方法：
- fetchPets()  — 调用 petsAPI.getAll() 获取可领养宠物列表
- handleSearchConfirm()  — 点击搜索按钮触发名称搜索
- handleClearAll()  — 一键清除所有筛选条件
- formatAgeMonths() / extractCity() / formatHealthStatus() / formatSizeDisplay() — 数据格式化工具函数

API 调用（frontend/src/api/pets.js 第14-21行）：
- petsAPI.getAll(params) → GET /pets/  — 带查询参数获取宠物列表

前端路由（frontend/src/App.js 第27行）：/pets → PetList''')

doc.add_heading('1.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端完整调用链】

1. 用户访问 /pets 页面 → React Router 渲染 PetList 组件
2. PetList.useEffect() 触发 → buildApiParams() 构建查询参数（adoption_status=available + 筛选条件）
3. 调用 petsAPI.getAll(params) → axios 发送 GET 请求到 /api/pets/?adoption_status=available&species=dog&...
4. Django URL 路由匹配 → urls.py DefaultRouter 分发到 PetProfileViewSet
5. PetProfileViewSet.get_permissions() → list 操作返回 AllowAny（公开访问）
6. PetProfileViewSet.get_queryset() → 链式调用的筛选逻辑：
   - super().get_queryset().select_related('rescue_case') 预加载关联数据
   - 按 species/gender/size_category/adoption_status/search/location/年龄范围 逐层过滤
   - is_public 过滤：非管理员只能看到公开宠物
7. PetProfileSerializer 序列化数据 → 附加 rescue_case_address, rescue_case_appearance, size_category_display
8. 返回 JSON 响应 → 前端 PetList 更新 pets 状态 → 渲染卡片列表
''')

doc.add_heading('1.3 源代码', level=2)

# 后端 views.py PetProfileViewSet
add_code_block(doc, open(r'd:\software\pet-connect-web\backend\pets\views.py', encoding='utf-8').read()[:4000],
               'backend/pets/views.py（第1-71行，PetProfileViewSet完整代码）')

# 后端 urls.py
add_code_block(doc, open(r'd:\software\pet-connect-web\backend\pets\urls.py', encoding='utf-8').read(),
               'backend/pets/urls.py（完整文件）')

# 后端 serializers.py PetProfileSerializer
serializers_code = open(r'd:\software\pet-connect-web\backend\pets\serializers.py', encoding='utf-8').read()
# 提取 PetProfileSerializer 部分（第1-28行）
add_code_block(doc, '\n'.join(serializers_code.split('\n')[:29]),
               'backend/pets/serializers.py（第1-28行，PetProfileSerializer）')

# 后端 models.py PetProfile
models_code = open(r'd:\software\pet-connect-web\backend\pets\models.py', encoding='utf-8').read()
add_code_block(doc, '\n'.join(models_code.split('\n')[:40]),
               'backend/pets/models.py（第1-39行，PetProfile 模型）')

# 前端 PetList.js
petlist_code = open(r'd:\software\pet-connect-web\frontend\src\pages\PetList.js', encoding='utf-8').read()
add_code_block(doc, petlist_code, 'frontend/src/pages/PetList.js（完整文件）')

# 前端 API
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\api\pets.js', encoding='utf-8').read(),
               'frontend/src/api/pets.js（petsAPI 和 uploadAPI）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\constants\site.js', encoding='utf-8').read(),
               'frontend/src/constants/site.js（全局常量定义）')


# ===== 1.2 宠物详情页浏览 =====
doc.add_heading('2 宠物详情页浏览功能', level=1)

doc.add_heading('2.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('文件：backend/pets/views.py 第19-71行 — PetProfileViewSet 类（同宠物列表，retrieve 操作复用同一个 ViewSet）')
doc.add_paragraph('''
复用 PetProfileViewSet：
- retrieve 操作（GET /pets/{id}/）：由 ModelViewSet 自动提供，返回单条宠物详情
- get_permissions() 对 retrieve 操作返回 AllowAny（公开访问）
- get_queryset() 中对 retrieve 操作同样应用 is_public 过滤逻辑

序列化器增强（backend/pets/serializers.py 第8-28行）：
PetProfileSerializer 添加了三个 SerializerMethodField：
- rescue_case_address → 关联救助案例的发现地址
- rescue_case_appearance → 关联救助案例的外观描述
- size_category_display → 体型的展示名称（中文）''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('文件：frontend/src/pages/PetDetail.js 第159-920行 — PetDetail 组件')
doc.add_paragraph('''
组件名：PetDetail (React 函数组件)

核心 Hook：
- useState: pet(当前宠物), allPets(全部可领养宠物), loading, error, username, adoptStep(1-3), applicationId, message, questionnaire, attachmentFile, submitting, submitError, submitSuccess, showAdoptForm
- useEffect: 根据 id 参数并行请求宠物详情和全部可领养宠物列表

主要方法：
- fetchData(): 并行调用 petsAPI.getById(id) 和 petsAPI.getAll()
- goToPrev()/goToNext(): 前后切换宠物（使用 allPets 列表支持左右箭头导航）
- handleStep1Submit(): 创建领养申请 → adoptAPI.create(pet_id, message)
- handleStep2Submit(): 提交问卷 → adoptAPI.submitQuestionnaire(applicationId, questionnaire)
- handleStep3Submit(): 上传附件 → uploadAPI.upload() + adoptAPI.addAttachment()

左侧展示：大图 + 前后箭头 + 照片下方信息网格（种类/性别/年龄/体型/地区/健康/救助编号）
右侧展示：问候语卡片（"你好！{username} 我是{name}" + "带我回家"按钮）
下方展示：捡拾详情描述、领养申请表单（三步向导）

前端路由：/pets/:id → PetDetail''')

doc.add_heading('2.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端完整调用链】

1. 用户在宠物列表点击卡片 → navigate(/pets/{id}) → React Router 渲染 PetDetail 组件
2. PetDetail.useEffect() 并行请求：
   - petsAPI.getById(id) → GET /api/pets/{id}/
   - petsAPI.getAll({adoption_status:'available'}) → GET /api/pets/?adoption_status=available
3. 后端 PetProfileViewSet.retrieve() → 返回序列化后的宠物详情（含关联案例地址等扩展字段）
4. getAll() 返回全部可领养宠物 → 用于前后切换导航
5. 前端渲染：大图 + 问候语 + 宠物信息网格 + 描述文字
6. 用户点击"带我回家" → 跳转至独立的领养表单页 /adopt/{petId}
   （或在页面内展开三步申请表单，调用 adoptAPI.create/submitQuestionnaire/addAttachment）
''')

doc.add_heading('2.3 源代码', level=2)

# PetProfileSerializer 完整
add_code_block(doc, '\n'.join(serializers_code.split('\n')),
               'backend/pets/serializers.py（完整文件，含 PetProfileSerializer）')

# PetDetail.js
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\PetDetail.js', encoding='utf-8').read(),
               'frontend/src/pages/PetDetail.js（完整文件）')


# ===== 1.3 领养问卷提交 =====
doc.add_heading('3 领养问卷提交功能', level=1)

doc.add_heading('3.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
涉及三个核心类：

1. AdoptApplicationViewSet（backend/pets/views.py 第74-132行）
   - perform_create(self, serializer): 第89-92行，创建申请时自动设置 applicant 并修改宠物状态为 pending
   - questionnaire(self, request, pk): 第99-112行，自定义 action (POST /adopt/applications/{id}/questionnaire/)
     验证申请人身份 → 检查问卷是否已存在 → 创建 AdoptQuestionnaire 记录
   - attachments(self, request, pk): 第114-132行，自定义 action (POST /adopt/applications/{id}/attachments/)
     验证申请人身份 → 创建 AdoptAttachment 记录

2. AdoptQuestionnaireSerializer（backend/pets/serializers.py 第74-78行）
   - 字段：application(只读), answers_json(JSONField), submitted_at(只读)

3. AdoptAttachmentSerializer（backend/pets/serializers.py 第81-90行）
   - validate_file_size(): 第87-89行，验证文件大小必须大于0

数据模型：
- AdoptQuestionnaire（backend/pets/models.py 第67-76行）：OneToOneField 关联 AdoptApplication，answers_json 存JSON
- AdoptAttachment（backend/pets/models.py 第78-87行）：ForeignKey 关联 AdoptApplication，多对一关系''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('文件：frontend/src/pages/AdoptApplication.js 第59-1203行 — AdoptApplication 组件')
doc.add_paragraph('''
组件名：AdoptApplication (React 函数组件)

步骤流程（3步向导）：
- 步骤1（填写信息）：10个字段表单（姓名/地址/电话/微信/性别/年龄/住宅类型/住房情况/户外区域/同住人同意/养宠经历/未来变化）
  → 验证通过后调用 adoptAPI.create() + adoptAPI.submitQuestionnaire()
- 步骤2（上传材料）：上传身份证 + 居住证明（支持jpg/png/pdf，限5MB）
  → 调用 uploadAPI.upload() + adoptAPI.addAttachment()
- 步骤3（完成）：显示成功页面，引导返回

核心方法：
- validateStep1(): 第130-162行，完整表单校验（必填、格式、联动校验）
- handleStep1Submit(): 第164-193行，创建申请 + 提交问卷
- handleStep2Submit(): 第225-263行，上传两个文件并关联到申请
- validateFile(): 第197-207行，文件类型/大小校验

前端路由：/adopt/:petId → AdoptApplication''')

doc.add_heading('3.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端完整调用链】

步骤1（填写问卷）：
1. 用户在 /pets/{id} 点击"带我回家" → navigate(/adopt/{petId})
2. AdoptApplication 组件加载 → petsAPI.getById(petId) 获取宠物基本信息
3. 用户填写10字段表单 → validateStep1() 校验
4. handleStep1Submit() 调用：
   a) adoptAPI.create({pet_id, message}) → POST /api/adopt/applications/
      → AdoptApplicationViewSet.create() → perform_create(): 设置 applicant, 更新宠物状态为 pending
   b) adoptAPI.submitQuestionnaire(appId, answers) → POST /api/adopt/applications/{id}/questionnaire/
      → AdoptApplicationViewSet.questionnaire(): 校验身份 → 创建 AdoptQuestionnaire

步骤2（上传材料）：
5. 用户选择身份证和居住证明文件
6. handleStep2Submit() 调用：
   a) uploadAPI.upload(idCardFile, 'adopt') → POST /api/uploads/ → 返回文件URL
   b) adoptAPI.addAttachment(appId, {file_type:'id_card', file_url, file_size})
      → POST /api/adopt/applications/{id}/attachments/
      → AdoptApplicationViewSet.attachments(): 校验身份 → 创建 AdoptAttachment
   c) 同理上传居住证明
7. 全部完成 → 显示成功页面
''')

doc.add_heading('3.3 源代码', level=2)

# 后端 views.py AdoptApplicationViewSet 部分
views_code = open(r'd:\software\pet-connect-web\backend\pets\views.py', encoding='utf-8').read()
add_code_block(doc, views_code,
               'backend/pets/views.py（完整文件）')

add_code_block(doc, '\n'.join(models_code.split('\n')[41:]),
               'backend/pets/models.py（第41-107行，AdoptApplication/AdoptQuestionnaire/AdoptAttachment/AdoptOfflineVerify 模型）')

add_code_block(doc, open(r'd:\software\pet-connect-web\backend\pets\adopt_urls.py', encoding='utf-8').read(),
               'backend/pets/adopt_urls.py（完整文件，领养申请路由配置）')

# 前端
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\AdoptApplication.js', encoding='utf-8').read(),
               'frontend/src/pages/AdoptApplication.js（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\api\adopt.js', encoding='utf-8').read(),
               'frontend/src/api/adopt.js（完整文件，领养相关API封装）')


# ===== 1.4 用户审核与核验状态查看 =====
doc.add_heading('4 用户审核与核验状态查看功能', level=1)

doc.add_heading('4.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
AdoptApplicationViewSet（backend/pets/views.py 第74-132行）：
- get_queryset(self): 第84-87行，my/retrieve 操作仅返回当前用户的申请
- my(self, request): 第94-97行，@action(detail=False, url_path='my')，GET /adopt/applications/my/
- get_permissions(): 第79-82行，my/retrieve 操作需认证

AdoptApplicationSerializer（backend/pets/serializers.py 第31-71行）：
- get_verify_status(self, obj): 第50-54行，通过 offline_verify 反向关联获取核验状态
- get_verify_note(self, obj): 第56-60行，获取核验备注
- validate(): 第62-71行，防止用户对同一宠物重复提交活跃申请''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
两个前端页面：

1. MyApplications.js（第48-455行）— 审核状态列表
   - 调用 adoptAPI.getMy() 获取当前用户的所有申请
   - getDisplayStatus(): 第6-32行，综合 online_status + verify_status 显示中文状态标签
   - 分页显示（每页10条），表格列：宠物名、状态标签、提交时间、详情按钮
   - 路由：/my-applications

2. ApplicationDetail.js（第42-369行）— 驳回/失败详情
   - 调用 adoptAPI.getById(id) 获取单条申请详情
   - getDetailType(): 第18-40行，判断是审核拒绝还是核验失败，提取原因和时间
   - 展示驳回原因卡片（红色主题，含审核/核验时间和原因）
   - 路由：/my-applications/:id''')

doc.add_heading('4.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端调用链】

我的申请列表：
1. 用户访问 /my-applications → MyApplications 组件加载
2. useEffect() → adoptAPI.getMy() → GET /api/adopt/applications/my/
3. 后端 AdoptApplicationViewSet.my() → get_queryset() 过滤 applicant=当前用户
4. AdoptApplicationSerializer 序列化（含 verify_status, verify_note 通过 SerializerMethodField 从关联表中提取）
5. 前端 getDisplayStatus() 合并 online_status 和 verify_status → 渲染状态标签

申请详情（驳回/失败）：
1. 用户在列表中点击"查看详情" → /my-applications/{id} → ApplicationDetail 组件加载
2. useEffect() → adoptAPI.getById(id) → GET /api/adopt/applications/{id}/
3. 后端 AdoptApplicationViewSet.retrieve() → get_queryset() 过滤 applicant=当前用户
4. 返回完整申请数据（含 audit_opinion, verify_note 等字段）
5. 前端 getDetailType() 判断详情类型 → 渲染驳回原因''')

doc.add_heading('4.3 源代码', level=2)

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\MyApplications.js', encoding='utf-8').read(),
               'frontend/src/pages/MyApplications.js（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\ApplicationDetail.js', encoding='utf-8').read(),
               'frontend/src/pages/ApplicationDetail.js（完整文件）')


# ===== 1.5 管理员审核与核验 =====
doc.add_heading('5 管理员审核与核验功能', level=1)

doc.add_heading('5.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
两个管理员专用 ViewSet：

1. AdminAdoptApplicationViewSet（backend/pets/views.py 第135-150行）
   - permission_classes = [IsAdminRole]  — 仅管理员可访问
   - serializer_class = AdoptApplicationAuditSerializer
   - update(self, request, pk): 第140-150行
     → 验证审核数据 → 调用 serializer.save() 执行审核逻辑 → 记录操作日志
   - URL：PUT/PATCH /adopt/applications/{id}/audit/

2. AdminOfflineVerifyViewSet（backend/pets/views.py 第153-185行）
   - permission_classes = [IsAdminRole]
   - update(self, request, pk): 第158-180行
     → 核验失败时校验原因必填 → 更新核验状态 → 联动更新申请状态和宠物状态
     · passed: application→approved, pet→adopted
     · failed: application→rejected, pet→available
   - create_for_application(): 第182-185行，为申请创建核验记录
   - URL：PUT/PATCH /adopt/offline-verify/{id}/

3. AdoptApplicationAuditSerializer（backend/pets/serializers.py 第100-125行）
   - validate(): 第105-109行，拒绝时必须填写驳回原因
   - update(): 第112-125行，执行审核的核心逻辑：
     · 批准 → 宠物状态设为 pending
     · 拒绝 → 宠物恢复为 available（仅在无其他活跃申请时）

权限类 IsAdminRole：检查 profile.role == 'admin' 或 is_superuser''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
管理员审核与核验功能的前端实现分布在以下组件中：

1. AdminManageBar 组件（frontend/src/components/AdminManageBar.js）
   - 在宠物列表和管理页面渲染"编辑/隐藏/删除"操作按钮
   - 仅管理员可见（通过 role === 'admin' 判断）

2. 管理后台页面
   - 路由 /admin/* 系列页面
   - 调用 adoptAPI.audit() 和 adoptAPI.offlineVerify() 进行审核和核验操作

API 封装（frontend/src/api/adopt.js 第10-11行）：
- audit(id, data) → PUT /adopt/applications/{id}/audit/
- offlineVerify(id, data) → PUT /adopt/offline-verify/{id}/''')

doc.add_heading('5.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【管理员审核调用链】

1. 管理员在后台查看申请列表 → 选择某条申请点击"审核"
2. 前端调用 adoptAPI.audit(id, {online_status, audit_opinion})
   → PUT /api/adopt/applications/{id}/audit/
3. Django URL 路由匹配 → AdminAdoptApplicationViewSet.update()
4. AdoptApplicationAuditSerializer 处理：
   a) validate(): 如果是拒绝操作，检查 audit_opinion 非空
   b) update():
      - 设置审核人、审核时间
      - 批准 → 宠物 adoption_status = 'pending'
      - 拒绝 → 宠物恢复为 'available'（需确认无其他活跃申请）
5. write_operation_log() 记录操作日志（模块=adopt, 动作=audit）
6. 返回更新后的申请数据

【管理员核验调用链】

1. 管理员在后台选择某条已批准的申请进行线下核验
2. 前端调用 adoptAPI.offlineVerify(id, {verify_status, verify_note})
   → PUT /api/adopt/offline-verify/{id}/
3. Django URL 路由匹配 → AdminOfflineVerifyViewSet.update()
4. 校验失败时必须有 reason → 更新核验记录
5. 联动更新：
   - passed → application.online_status='approved' + pet.adoption_status='adopted'
   - failed → application.online_status='rejected' + pet.adoption_status='available'
6. 返回核验结果

【状态联动规则总结】
- 提交申请 → pet: available → pending
- 审核批准 → pet: pending (不变)
- 审核拒绝 → pet: pending → available（无其他活跃申请时）
- 核验通过 → application: approved + pet: adopted
- 核验失败 → application: rejected + pet: available
''')

doc.add_heading('5.3 源代码', level=2)

add_code_block(doc, views_code,
               'backend/pets/views.py（完整文件，含 AdminAdoptApplicationViewSet 和 AdminOfflineVerifyViewSet）')

add_code_block(doc, '\n'.join(serializers_code.split('\n')[99:]),
               'backend/pets/serializers.py（第99-126行，AdoptApplicationAuditSerializer）')

add_code_block(doc, '\n'.join(models_code.split('\n')[88:]),
               'backend/pets/models.py（第88-107行，AdoptOfflineVerify 模型）')


# =============================================
# 模块二：救助追踪模块
# =============================================
doc.add_heading('模块二：救助追踪模块', level=1)

# ===== 2.1 救助列表浏览 =====
doc.add_heading('6 救助列表浏览功能', level=1)

doc.add_heading('6.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
RescueCaseViewSet（backend/rescue/views.py 第21-201行）：
- queryset: 使用 select_related('reporter') 和 prefetch_related('status_logs') 优化查询
- pagination_class = RescueCasePagination（第15-18行，每页10条，最大50条）
- get_queryset(self): 第26-54行，支持多维度过滤：
  · my=true → 我上报的救助案例
  · helped=true → 我响应的救助案例
  · status → 按状态过滤（支持逗号分隔多状态）
  · rescue_no → 按救助编号精确搜索
  · exclude_status → 排除指定状态
- get_permissions(self): 第56-61行，list/retrieve 公开访问

RescueCaseSerializer（backend/rescue/serializers.py 第39-69行）：
- 嵌套 reporter(UserSerializer), helpers(多用户), status_logs(状态日志列表)
- 自定义 CoordinateField 用于经纬度字段（自动四舍五入到6位小数）
- 验证：discover_address / nickname / contact 非空校验''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
RescueList.js（第50-212行）— 救助列表组件
- 调用 rescueAPI.getMyCases() 获取救助案例（排除 awaiting_adoption 状态）
- 使用 localStorage 追踪用户已"救助"的记录，已响应过的会被过滤
- buildDescription(): 第27-48行，根据案例字段构建中文描述文本
- 卡片式布局：照片 + 描述 + 发布时间 + "救助"按钮
- 三个功能入口：我的救助、上报、查询

API 封装（frontend/src/api/pets.js 第23-33行）：
- rescueAPI.getAll(params) → GET /rescue/cases/
- rescueAPI.getMyCases() → GET /rescue/cases/?my=true&exclude_status=awaiting_adoption
- rescueAPI.help(id) → POST /rescue/cases/{id}/help/

前端路由：/rescue → RescueList''')

doc.add_heading('6.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端调用链】

1. 用户访问 /rescue 页面 → RescueList 组件加载
2. useEffect() → rescueAPI.getMyCases()
   → GET /api/rescue/cases/?my=true&exclude_status=awaiting_adoption
3. 后端 RescueCaseViewSet.get_queryset():
   - my=true → 过滤 reporter=当前用户
   - exclude_status=awaiting_adoption → 排除待领养状态
4. 前端加载 localStorage 中已救助列表 → 过滤已响应的案例
5. 渲染案例卡片 → 用户点击"救助"按钮
6. handleHelp() → rescueAPI.help(id) → POST /api/rescue/cases/{id}/help/
7. 后端 RescueCaseViewSet.help():
   - 检查是否已响应 → 添加 helpers 关联
   - 首次响应 → 记录 help_date，状态从 pending_rescue 推进到 in_medical
   - 记录 RescueStatusLog → 记录操作日志
8. 前端更新 localStorage → 从列表中移除该案例
''')

doc.add_heading('6.3 源代码', level=2)

# 后端 rescue
rescue_views_code = open(r'd:\software\pet-connect-web\backend\rescue\views.py', encoding='utf-8').read()
add_code_block(doc, rescue_views_code[:3800],
               'backend/rescue/views.py（第1-104行，RescueCaseViewSet 列表和帮助功能）')

rescue_models_code = open(r'd:\software\pet-connect-web\backend\rescue\models.py', encoding='utf-8').read()
add_code_block(doc, rescue_models_code[:60],
               'backend/rescue/models.py（第1-53行，RescueCase 模型完整定义）')

add_code_block(doc, open(r'd:\software\pet-connect-web\backend\rescue\serializers.py', encoding='utf-8').read(),
               'backend/rescue/serializers.py（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\backend\rescue\urls.py', encoding='utf-8').read(),
               'backend/rescue/urls.py（完整文件）')

# 前端
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\RescueList.js', encoding='utf-8').read(),
               'frontend/src/pages/RescueList.js（完整文件）')


# ===== 2.2 救助上报 =====
doc.add_heading('7 救助上报功能', level=1)

doc.add_heading('7.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
RescueCaseViewSet（backend/rescue/views.py）：
- perform_create(self, serializer): 第63-71行
  → 自动生成救助编号（调用 generate_rescue_no()）
  → 保存案例（关联当前用户为 reporter）
  → 创建初始 RescueStatusLog（from_status=None, to_status=pending_rescue, remark='被发现上报'）

generate_rescue_no()（backend/rescue/serializers.py 第81-87行）：
  → 格式：RES + 年月日 + 3位流水号（如 RES20260623001）
  → 按日期分组，每天从001开始递增

RescueCaseSerializer 字段校验：
- validate_discover_address(): 第51-55行，发现地点不能为空
- validate_nickname(): 第57-61行，昵称不能为空
- validate_contact(): 第63-69行，联系方式不能为空且长度≥5''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
RescueReport.js（第32-626行）— 救助上报表单组件

核心功能：
- 表单字段：昵称、联系方式、发现地址、照片上传、体型描述、健康状况、是否受伤、是否怕人
- GPS 自动定位：handleUseCurrentLocation() 第66-111行
  → 调用浏览器 navigator.geolocation.getCurrentPosition()
  → 通过高德地图逆地理编码 API 将坐标转为可读地址
- 照片上传：handlePhotoSelect() 第114-148行
  → 逐文件调用 uploadAPI.upload(f, 'rescue')
  → 支持多文件、格式校验（jpg/png/pdf）、大小限制（5MB）
- 表单校验：validate() 第161-175行，所有必填项校验
- 提交：handleSubmit() 第179-210行
  → 构建 payload（含经纬度坐标、格式化的 appearance 外观描述文本）
  → 调用 rescueAPI.create(payload)
- 成功后显示救助编号

前端路由：/rescue/report → RescueReport''')

doc.add_heading('7.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端调用链】

1. 用户在救助列表点击"上报" → /rescue/report → RescueReport 组件
2. 用户可选"定位"按钮 → 浏览器GPS定位 → 高德API逆地理编码 → 自动填充地址和坐标
3. 用户选择照片 → 逐文件上传 uploadAPI.upload(file, 'rescue')
   → POST /api/uploads/ → 后端保存文件到 rescue 子目录 → 返回文件URL
4. 用户填写全部必填字段 → validate() 校验通过
5. handleSubmit() → rescueAPI.create(payload) → POST /api/rescue/cases/
6. 后端 RescueCaseViewSet.perform_create():
   a) generate_rescue_no() 生成唯一救助编号
   b) serializer.save(reporter=request.user, rescue_no=...)
   c) RescueStatusLog.objects.create(...) 记录初始状态
7. 返回创建结果（含 rescue_no）→ 前端显示成功页和救助编号
''')

doc.add_heading('7.3 源代码', level=2)

add_code_block(doc, rescue_views_code,
               'backend/rescue/views.py（完整文件，含 perform_create 和所有 action）')

add_code_block(doc, rescue_models_code,
               'backend/rescue/models.py（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\RescueReport.js', encoding='utf-8').read(),
               'frontend/src/pages/RescueReport.js（完整文件）')


# ===== 2.3 我的救助 =====
doc.add_heading('8 我的救助功能', level=1)

doc.add_heading('8.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
RescueCaseViewSet.get_queryset()（第26-54行）：
- helped=true 查询参数 → 过滤 helpers 包含当前用户的救助案例
- 支持 status 参数按状态筛选
- 返回分页结果（RescueCasePagination，每页10条）''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
MyRescues.js（第23-254行）— 我的救助列表组件

核心功能：
- 调用 rescueAPI.getHelpedCases(params) 获取当前用户响应的救助案例
- 按状态筛选（下拉框）+ 分页（上一页/下一页）
- 表格展示：动物缩略图、救助编号、状态标签、救助日期、操作按钮
- showUpdateBtn() 第74行：非 rescued/abandoned 状态才显示操作按钮
- 操作按钮：更新状态（→ /my-rescues/{id}/update-status）、填写记录（→ /my-rescues/{id}/stage-record）

API 封装：
- rescueAPI.getHelpedCases(params) → GET /rescue/cases/?helped=true&page=N&status=X

前端路由：/my-rescues → MyRescues''')

doc.add_heading('8.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【前端 → 后端调用链】

1. 用户点击"我的救助" → /my-rescues → MyRescues 组件加载
2. useEffect() → rescueAPI.getHelpedCases({page, status})
   → GET /api/rescue/cases/?helped=true&page=N&status=X
3. 后端 RescueCaseViewSet.get_queryset():
   - helped=true → 过滤 helpers 包含当前用户
   - status=X → 过滤 current_status=X
   - 分页返回
4. 前端渲染表格 → 用户可选择筛选条件或翻页
5. 用户点击"更新状态" → /my-rescues/{id}/update-status
6. 用户点击"填写记录" → /my-rescues/{id}/stage-record
''')

doc.add_heading('8.3 源代码', level=2)

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\MyRescues.js', encoding='utf-8').read(),
               'frontend/src/pages/MyRescues.js（完整文件）')


# ===== 2.4 救助查询与详情 =====
doc.add_heading('9 救助查询与详情功能', level=1)

doc.add_heading('9.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
RescueCaseViewSet：
- get_queryset() 中 rescue_no 过滤（第46-48行）：rescue_no__iexact 精确不区分大小写匹配
- retrieve 操作：ModelViewSet 自动支持 GET /rescue/cases/{id}/
- stage_records action（第166-201行）：GET 列出 + POST 创建阶段记录

RescueStageRecordSerializer（serializers.py 第72-78行）：
- 嵌套 operator(UserSerializer)
- content 必填，rescue_case/operator/created_at 只读''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
两个前端组件：

1. RescueSearch.js（第15-125行）— 救助查询
   - 单个输入框输入救助编号 → 回车或点击"查询"
   - 调用 rescueAPI.getAll({rescue_no}) → GET /rescue/cases/?rescue_no=XXX
   - 结果列表：编号 + 状态标签（点击跳转到详情）
   - 路由：/rescue/search

2. RescueDetail.js（第32-296行）— 救助详情
   - 并行加载案例数据和阶段记录
   - 状态流转路径展示：已到达的状态按序显示为可点击标签
   - 时间线视图：合并 status_logs 和 stage_records，按时间分组到各状态下
   - 每个时间线条目显示：时间、操作人、内容
   - 图标区分：状态变更（交换图标）+ 阶段记录（剪贴板图标）
   - 状态元信息：悬停显示含义和预计停留时长
   - 路由：/rescue/:id''')

doc.add_heading('9.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【救助查询调用链】

1. 用户在 /rescue/search 输入救助编号 → handleSearch()
2. rescueAPI.getAll({rescue_no: trimmed}) → GET /api/rescue/cases/?rescue_no=XXX
3. 后端 RescueCaseViewSet.get_queryset() → qs.filter(rescue_no__iexact=rescue_no)
4. 返回匹配的案例列表 → 前端渲染结果

【救助详情调用链】

1. 用户在搜索结果点击某条 → navigate(/rescue/{id}) → RescueDetail 组件
2. useEffect() 并行请求：
   - rescueAPI.getById(id) → GET /api/rescue/cases/{id}/
   - rescueAPI.getStageRecords(id) → GET /api/rescue/cases/{id}/stage-records/
3. 后端返回案例详情（含嵌套的 reporter/helpers/status_logs）和阶段记录列表
4. 前端 useMemo 计算：
   - reachedStatuses: 从 status_logs 提取已到达的状态列表
   - timelineByStatus: 将 status_logs 和 stage_records 合并，按状态和时间分组排序
   - currentIdx: 当前状态在流转路径中的位置
5. 渲染状态流转路径 + 可切换的时间线视图
''')

doc.add_heading('9.3 源代码', level=2)

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\RescueSearch.js', encoding='utf-8').read(),
               'frontend/src/pages/RescueSearch.js（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\RescueDetail.js', encoding='utf-8').read(),
               'frontend/src/pages/RescueDetail.js（完整文件）')


# ===== 2.5 状态更新 =====
doc.add_heading('10 状态更新功能', level=1)

doc.add_heading('10.1 实现类（类保护的属性与方法）', level=2)

doc.add_heading('后端实现', level=3)
doc.add_paragraph('''
RescueCaseViewSet 中的状态更新相关方法：

1. help action（第74-104行）— 响应救助时的自动状态推进
   - POST /rescue/cases/{id}/help/
   - 用户加入 helpers → 若为首次响应且状态为 pending_rescue → 自动推进到 in_medical
   - 记录 help_date 和 RescueStatusLog

2. update_status action（第113-163行）— 手动状态更新
   - PATCH /rescue/cases/{id}/status/
   - 权限：仅 helpers 成员或管理员可操作
   - 三重校验：
     a) 状态值合法性（必须在 STATUS_CHOICES 中）
     b) 备注必填（诊断结果、治疗方案等）
     c) 状态流转规则（只能按 STATUS_FLOW 顺序推进）
   - STATUS_FLOW 流转规则（第107-111行）：
     · in_medical → recovering → awaiting_adoption → rescued
     · pending_rescue 和 abandoned 不允许手动变更
   - 成功后记录 RescueStatusLog 和操作日志

3. stage_records action（第166-201行）— 阶段记录
   - GET /rescue/cases/{id}/stage-records/ — 列出所有阶段记录
   - POST — 创建新记录（需 helpers 或管理员权限）
   - 内容必填校验

4. RescueStatusLog 模型（rescue/models.py 第55-65行）：
   - 记录每次状态变更的 from_status → to_status
   - 关联操作人(operator)和备注(remark)

5. RescueStageRecord 模型（rescue/models.py 第68-77行）：
   - 救助各阶段的详细记录，可多次填写
   - 关联救助案例、内容、操作人''')

doc.add_heading('前端实现', level=3)
doc.add_paragraph('''
两个前端组件：

1. UpdateRescueStatus.js（第23-285行）— 状态更新页面
   - 加载案例数据 → 显示当前状态标签
   - STATUS_NEXT 流转规则（第17-21行）：与后端 STATUS_FLOW 对应
   - 点击"切换状态" → 显示新状态预览 + 备注输入框
   - 备注校验（非空）→ 调用 rescueAPI.updateStatus(id, {current_status, remark})
   - 成功后显示新状态标签
   - 路由：/my-rescues/:id/update-status

2. StageRecord.js（第12-241行）— 阶段记录填写页面
   - 加载案例数据和已有记录（并行请求）
   - 文本区域输入当前阶段记录内容
   - 校验非空 → 调用 rescueAPI.addStageRecord(id, {content})
   - 成功后追加到记录列表顶部
   - 历史记录列表：按时间倒序展示，含操作人和时间
   - 路由：/my-rescues/:id/stage-record''')

doc.add_heading('10.2 各个类及方法的调用关系', level=2)
doc.add_paragraph('''
【状态更新调用链】

1. 用户在"我的救助"列表点击"更新状态" → /my-rescues/{id}/update-status
2. UpdateRescueStatus 组件加载 → rescueAPI.getById(id) 获取当前状态
3. 用户点击"切换状态" → 前端 STATUS_NEXT 匹配下一个状态
4. 用户填写备注 → 点击"确认更新"
5. rescueAPI.updateStatus(id, {current_status, remark})
   → PATCH /api/rescue/cases/{id}/status/
6. 后端 RescueCaseViewSet.update_status():
   a) 权限检查：is_helper 或 is_admin
   b) 状态合法性校验：new_status in STATUS_CHOICES
   c) 备注非空校验
   d) 流转规则校验：STATUS_FLOW[old_status] == new_status
   e) 更新 case.current_status
   f) 创建 RescueStatusLog
   g) write_operation_log() 记录审计日志
7. 返回更新后的案例数据 → 前端显示成功页

【阶段记录调用链】

1. 用户在"我的救助"列表点击"填写记录" → /my-rescues/{id}/stage-record
2. StageRecord 组件并行加载案例 + 已有记录
3. 用户填写内容 → 点击"提交记录"
4. rescueAPI.addStageRecord(id, {content}) → POST /api/rescue/cases/{id}/stage-records/
5. 后端 RescueCaseViewSet.stage_records(POST):
   a) 权限检查：is_helper 或 is_admin
   b) 内容非空校验
   c) RescueStageRecord.objects.create(...)
6. 返回新记录 → 前端插入列表顶部
''')

doc.add_heading('10.3 源代码', level=2)

# 后端完整 rescue views.py（含 update_status 和 stage_records action）
add_code_block(doc, rescue_views_code,
               'backend/rescue/views.py（完整文件，重点关注第106-201行的 update_status 和 stage_records）')

add_code_block(doc, rescue_models_code,
               'backend/rescue/models.py（完整文件）')

# 前端状态更新和阶段记录
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\UpdateRescueStatus.js', encoding='utf-8').read(),
               'frontend/src/pages/UpdateRescueStatus.js（完整文件）')

add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\pages\StageRecord.js', encoding='utf-8').read(),
               'frontend/src/pages/StageRecord.js（完整文件）')


# ===== 附录：前端路由配置 =====
doc.add_heading('附录：前端路由配置', level=1)
add_code_block(doc, open(r'd:\software\pet-connect-web\frontend\src\App.js', encoding='utf-8').read(),
               'frontend/src/App.js（路由配置，完整文件）')


# ===== 保存文档 =====
output_path = r'd:\software\pet-connect-web\docs\宠物领养与救助追踪模块编码文档.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')

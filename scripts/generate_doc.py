"""
生成宠物领养和救助追踪模块编码文档（Word格式）
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# 项目根目录
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

doc = Document()

# ===== 设置默认字体 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
# 中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        heading_font.size = Pt(18)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 0, 0)
    elif i == 2:
        heading_font.size = Pt(15)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 0, 0)
    elif i == 3:
        heading_font.size = Pt(13)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 0, 0)

# ===== 辅助函数 =====
def add_code_block(doc, file_path_label, code_text, language="python"):
    """添加带文件路径标注和关键注释的代码块"""
    # 文件路径标题
    p = doc.add_paragraph()
    run = p.add_run(f'📄 文件路径：{file_path_label}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0, 100, 0)

    # 代码内容
    p = doc.add_paragraph()
    # 使用更小的字体以容纳完整代码
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # 设置段落格式
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1

    # 分隔线
    doc.add_paragraph('─' * 80)


def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    doc.add_paragraph()  # 表后空行
    return table


def read_file(rel_path):
    """读取文件内容，处理编码"""
    full_path = os.path.join(BASE_DIR, rel_path)
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(full_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, FileNotFoundError):
            continue
    return f'[无法读取文件: {full_path}]'

# ======================================================================
#                          封面
# ======================================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('宠物领养与救助追踪模块\n编码文档')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PetConnect（暖爪救助）')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('技术栈：Python 3.11 + Django REST Framework + React\n文档生成日期：2026年6月23日')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(130, 130, 130)

doc.add_page_break()

# ======================================================================
#                          目录页
# ======================================================================
doc.add_heading('目录', level=1)
toc_items = [
    '1  宠物领养功能',
    '    1.1  实现类（类保护的属性与方法）',
    '    1.2  各个类及方法的调用关系',
    '    1.3  源代码',
    '2  救助追踪功能',
    '    2.1  实现类（类保护的属性与方法）',
    '    2.2  各个类及方法的调用关系',
    '    2.3  源代码',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ======================================================================
#              1  宠物领养功能
# ======================================================================
doc.add_heading('1  宠物领养功能', level=1)

doc.add_paragraph(
    '宠物领养模块负责管理可领养宠物的档案信息、处理用户的领养申请流程、'
    '收集领养问卷、存储证明材料，以及管理员审核和线下核验的全流程。'
    '该模块由后端 Django 应用 pets 和前端 React 页面共同组成。'
)

# ---- 1.1 实现类 ----
doc.add_heading('1.1  实现类（类保护的属性与方法）', level=2)

doc.add_heading('1.1.1  后端模型类（Models）', level=3)

doc.add_paragraph(
    '以下为宠物领养模块的5个数据模型类，均继承自 django.db.models.Model。'
    'Django 模型字段默认使用 public 访问（Python 无严格访问控制），'
    '通过 Meta 内部类配置表名和排序。'
)

add_table(doc,
    ['类名', '继承自', '主要属性/字段', '访问修饰符', '文件路径'],
    [
        ['PetProfile',
         'models.Model',
         'ADOPTION_STATUS_CHOICES (类常量)\n'
         'SIZE_CATEGORY_CHOICES (类常量)\n'
         'rescue_case: ForeignKey→RescueCase\n'
         'name: CharField(100)\n'
         'species: CharField(50)\n'
         'breed: CharField(100)\n'
         'age_months: IntegerField\n'
         'gender: CharField(10)\n'
         'size_category: CharField(10)\n'
         'health_status: CharField(100)\n'
         'description: TextField\n'
         'photo_url: CharField(500)\n'
         'adoption_status: CharField(20)\n'
         'is_public: BooleanField\n'
         'created_at: DateTimeField\n'
         'updated_at: DateTimeField\n'
         'Meta: db_table="pet_profile"',
         'public (类常量)\npublic (所有字段)',
         'backend/pets/models.py'],

        ['AdoptApplication',
         'models.Model',
         'ONLINE_STATUS_CHOICES (类常量)\n'
         'applicant: ForeignKey→User (CASCADE)\n'
         'pet: ForeignKey→PetProfile (RESTRICT)\n'
         'online_status: CharField(20)\n'
         'audit_opinion: TextField\n'
         'auditor: ForeignKey→User (SET_NULL)\n'
         'audited_at: DateTimeField\n'
         'message: TextField\n'
         'created_at: DateTimeField\n'
         'updated_at: DateTimeField\n'
         'Meta: db_table="adopt_application"',
         'public',
         'backend/pets/models.py'],

        ['AdoptQuestionnaire',
         'models.Model',
         'application: OneToOneField→AdoptApplication (CASCADE)\n'
         'answers_json: JSONField\n'
         'submitted_at: DateTimeField\n'
         'Meta: db_table="adopt_questionnaire"',
         'public',
         'backend/pets/models.py'],

        ['AdoptAttachment',
         'models.Model',
         'application: ForeignKey→AdoptApplication (CASCADE)\n'
         'file_type: CharField(30)\n'
         'file_url: CharField(500)\n'
         'file_size: IntegerField\n'
         'uploaded_at: DateTimeField\n'
         'Meta: db_table="adopt_attachment"',
         'public',
         'backend/pets/models.py'],

        ['AdoptOfflineVerify',
         'models.Model',
         'VERIFY_STATUS_CHOICES (类常量)\n'
         'application: OneToOneField→AdoptApplication (CASCADE)\n'
         'verify_status: CharField(20)\n'
         'verify_note: TextField\n'
         'verifier: ForeignKey→User (SET_NULL)\n'
         'verified_at: DateTimeField\n'
         'Meta: db_table="adopt_offline_verify"',
         'public',
         'backend/pets/models.py'],
    ]
)

doc.add_heading('1.1.2  后端序列化器类（Serializers）', level=3)

add_table(doc,
    ['类名', '继承自', '主要属性/方法', '访问修饰符', '文件路径'],
    [
        ['PetProfileSerializer',
         'serializers.ModelSerializer',
         'rescue_case_address: SerializerMethodField\n'
         'rescue_case_appearance: SerializerMethodField\n'
         'size_category_display: SerializerMethodField\n'
         'get_rescue_case_address(self, obj)\n'
         'get_rescue_case_appearance(self, obj)\n'
         'get_size_category_display(self, obj)\n'
         'Meta: model=PetProfile, fields="__all__"',
         'public',
         'backend/pets/serializers.py'],

        ['AdoptApplicationSerializer',
         'serializers.ModelSerializer',
         'applicant: UserSerializer(read_only)\n'
         'pet: PetProfileSerializer(read_only)\n'
         'pet_id: PrimaryKeyRelatedField(write_only)\n'
         'verify_status: SerializerMethodField\n'
         'verify_note: SerializerMethodField\n'
         'get_verify_status(self, obj)\n'
         'get_verify_note(self, obj)\n'
         'validate(self, attrs) — 防重复申请校验',
         'public',
         'backend/pets/serializers.py'],

        ['AdoptQuestionnaireSerializer',
         'serializers.ModelSerializer',
         'Meta: model=AdoptQuestionnaire\n'
         'fields=[id,application,answers_json,submitted_at]\n'
         'read_only_fields=[application,submitted_at]',
         'public',
         'backend/pets/serializers.py'],

        ['AdoptAttachmentSerializer',
         'serializers.ModelSerializer',
         'validate_file_size(self, value) — 文件大小>0校验\n'
         'Meta: model=AdoptAttachment\n'
         'read_only_fields=[application,uploaded_at]',
         'public',
         'backend/pets/serializers.py'],

        ['AdoptOfflineVerifySerializer',
         'serializers.ModelSerializer',
         'Meta: model=AdoptOfflineVerify\n'
         'read_only_fields=[application,verifier,verified_at]',
         'public',
         'backend/pets/serializers.py'],

        ['AdoptApplicationAuditSerializer',
         'serializers.ModelSerializer',
         'validate(self, attrs) — 拒绝必填原因\n'
         'update(self, instance, validated_data)\n'
         '  — 审核后联动更新PetProfile.adoption_status',
         'public',
         'backend/pets/serializers.py'],
    ]
)

doc.add_heading('1.1.3  后端视图类（Views）', level=3)

add_table(doc,
    ['类名', '继承自', '主要方法', '权限控制', '文件路径'],
    [
        ['PetProfileViewSet',
         'viewsets.ModelViewSet',
         'get_permissions(self)\n'
         '  list/retrieve: AllowAny\n'
         '  my_pets: IsAuthenticated\n'
         '  其他: IsAdminRole\n'
         'get_queryset(self) — 支持多条件筛选\n'
         'my_pets(self, request) — @action GET /my/',
         '分层权限',
         'backend/pets/views.py'],

        ['AdoptApplicationViewSet',
         'viewsets.ModelViewSet',
         'http_method_names — 禁用了DELETE\n'
         'get_permissions(self)\n'
         '  create/my/retrieve等: IsAuthenticated\n'
         '  其他: IsAdminRole\n'
         'get_queryset(self) — my/retrieve仅返回自己的\n'
         'perform_create(self, serializer) — 自动关联用户\n'
         'my(self, request) — @action GET /my/\n'
         'questionnaire(self, request, pk) — POST 问卷\n'
         'attachments(self, request, pk) — POST 附件',
         '分层权限',
         'backend/pets/views.py'],

        ['AdminAdoptApplicationViewSet',
         'viewsets.GenericViewSet',
         'permission_classes = [IsAdminRole]\n'
         'update(self, request, pk) — 审核申请\n'
         '  → 调用write_operation_log记录日志',
         '[IsAdminRole]',
         'backend/pets/views.py'],

        ['AdminOfflineVerifyViewSet',
         'viewsets.GenericViewSet',
         'permission_classes = [IsAdminRole]\n'
         'update(self, request, pk) — 线下核验\n'
         '  核验通过→申请approved+宠物adopted\n'
         '  核验失败→申请rejected+宠物available\n'
         'create_for_application(self, request, app_id)',
         '[IsAdminRole]',
         'backend/pets/views.py'],
    ]
)

doc.add_heading('1.1.4  权限类（Permissions）', level=3)

add_table(doc,
    ['类名/函数名', '继承自', '主要方法', '访问修饰符', '文件路径'],
    [
        ['user_is_admin(user)',
         '模块级函数',
         '检查 user.is_superuser/is_staff\n或 profile.role == "admin"',
         'public（模块导出）',
         'backend/common/permissions.py'],

        ['IsAdminRole',
         'rest_framework.permissions.BasePermission',
         'has_permission(self, request, view)\n→ 调用 user_is_admin(request.user)',
         'public',
         'backend/common/permissions.py'],

        ['IsActiveUser',
         'rest_framework.permissions.BasePermission',
         'has_permission(self, request, view)\n→ 检查 profile.status == 0（未封禁）',
         'public',
         'backend/common/permissions.py'],
    ]
)

doc.add_heading('1.1.5  前端组件与API类', level=3)

add_table(doc,
    ['组件/模块名', '类型', '主要属性/方法', '路由', '文件路径'],
    [
        ['adoptAPI',
         'API模块（对象）',
         'getAll(), getById(id), create(data),\n'
         'getMy(), submitQuestionnaire(id,json),\n'
         'addAttachment(id,data), audit(id,data),\n'
         'offlineVerify(id,data)',
         '—',
         'frontend/src/api/adopt.js'],

        ['petsAPI',
         'API模块（对象）',
         'getAll(params), getById(id), create(data),\n'
         'update(id,data), delete(id), getMyPets()',
         '—',
         'frontend/src/api/pets.js'],

        ['uploadAPI',
         'API模块（对象）',
         'upload(file, subdir) → FormData上传',
         '—',
         'frontend/src/api/pets.js'],

        ['PetList',
         'React Function Component',
         'State: pets,loading,error,search,speciesFilter,\n'
         '  genderFilter,ageRangeFilter,locationFilter\n'
         'Methods: buildApiParams(), handleSearchConfirm(),\n'
         '  handleClearAll(), extractCity(),\n'
         '  formatHealthStatus(), formatAgeMonths(),\n'
         '  formatSizeDisplay()',
         '/pets',
         'frontend/src/pages/PetList.js'],

        ['PetDetail',
         'React Function Component',
         'State: pet,allPets,loading,error,adoptStep,\n'
         '  applicationId,message,questionnaire,\n'
         '  attachmentFile,fileType,submitting\n'
         'Methods: goToPrev(),goToNext(),requireAuth(),\n'
         '  handleStep1Submit(),handleStep2Submit(),\n'
         '  handleStep3Submit()',
         '/pets/:id',
         'frontend/src/pages/PetDetail.js'],

        ['AdoptApplication',
         'React Function Component',
         'State: pet,loading,error,step,applicationId,\n'
         '  form(16字段),errors,idCardFile,\n'
         '  housingProofFile,fileErrors\n'
         'Methods: handleChange(),validateStep1(),\n'
         '  handleStep1Submit(),handleStep2Submit(),\n'
         '  validateFile(),handleIdCardSelect()',
         '/adopt/:petId',
         'frontend/src/pages/AdoptApplication.js'],

        ['MyApplications',
         'React Function Component',
         'State: applications,loading,error,currentPage\n'
         'Methods: getDisplayStatus(),formatDateTime(),\n'
         '  renderPagination(),goToPage()\n'
         '常量: PAGE_SIZE=10',
         '/my-applications',
         'frontend/src/pages/MyApplications.js'],

        ['ApplicationDetail',
         'React Function Component',
         'State: application,loading,error\n'
         'Methods: getDetailType(),formatDateTime()\n'
         '— 显示审核拒绝/核验失败的详情原因',
         '/my-applications/:id',
         'frontend/src/pages/ApplicationDetail.js'],

        ['AddPet',
         'React Function Component',
         'State: authorized,authChecking,rescueCases,\n'
         '  form(10字段),photoUploading,loading\n'
         'Methods: handleChange(),handlePhotoChange(),\n'
         '  handleSubmit() — AI品种识别+文案生成',
         '/add-pet',
         'frontend/src/pages/AddPet.js'],

        ['AdminManageBar',
         'React Function Component',
         'Props: onEdit,onHide,onDelete\n'
         '— 通过useManageMode()检查canManage',
         '—（通用组件）',
         'frontend/src/components/AdminManageBar.js'],

        ['ProtectedRoute',
         'React Function Component',
         '— 检查localStorage token\n未登录→跳转/login',
         '—（通用组件）',
         'frontend/src/components/ProtectedRoute.js'],
    ]
)

doc.add_page_break()

# ---- 1.2 调用关系 ----
doc.add_heading('1.2  各个类及方法的调用关系', level=2)

doc.add_heading('1.2.1  后端调用链', level=3)

doc.add_paragraph(
    '【前端请求 → URL路由 → 视图 → 序列化器 → 模型 → 数据库】\n\n'
    '一、宠物档案管理流程：\n'
    '  1. 前端 PetList/PetDetail 组件 → petsAPI.getAll()/getById() → GET /api/pets/\n'
    '     → urls.py (DefaultRouter) → PetProfileViewSet.list()/retrieve()\n'
    '     → get_queryset() 多条件筛选（species/gender/age/search/location/size_category/is_public）\n'
    '     → PetProfileSerializer 序列化（含 rescue_case_address/rescue_case_appearance 跨表查询）\n'
    '     → Response 返回 JSON\n\n'
    '  2. 前端 AddPet 组件 → petsAPI.create(data) → POST /api/pets/\n'
    '     → PetProfileViewSet.create() → PetProfileSerializer(data).is_valid()\n'
    '     → PetProfile.objects.create(...) → Response 201\n'
    '     注：AddPet 需验证admin角色 → authAPI.getProfile() → profile.role === "admin"\n\n'
    '  3. 前端 AdminDashboard → petsAPI.update()/delete() → PATCH/DELETE /api/pets/{id}/\n'
    '     → PetProfileViewSet.partial_update()/destroy() → [IsAdminRole]权限校验\n\n'
    '二、领养申请流程：\n'
    '  4. 前端 AdoptApplication/PetDetail 组件 → adoptAPI.create(data) → POST /api/adopt/applications/\n'
    '     → AdoptApplicationViewSet.create() → perform_create()\n'
    '       → serializer.save(applicant=request.user)\n'
    '       → 联动更新 pet.adoption_status = "pending"（防止重复申请）\n'
    '     → AdoptApplicationSerializer.validate() 校验重复活跃申请\n\n'
    '  5. adoptAPI.submitQuestionnaire(id, json) → POST /api/adopt/applications/{id}/questionnaire/\n'
    '     → AdoptApplicationViewSet.questionnaire()\n'
    '       → 校验 applicant == request.user（仅申请人本人可提交）\n'
    '       → 校验 hasattr(app, "questionnaire") 防止重复提交\n'
    '       → AdoptQuestionnaireSerializer(data).is_valid()\n'
    '       → AdoptQuestionnaire.objects.create(application=app, answers_json=...) → Response 201\n\n'
    '  6. adoptAPI.addAttachment(id, data) → POST /api/adopt/applications/{id}/attachments/\n'
    '     → AdoptApplicationViewSet.attachments()\n'
    '       → 校验 applicant == request.user\n'
    '       → AdoptAttachmentSerializer(data).is_valid()\n'
    '       → AdoptAttachment.objects.create(...) → Response 201\n\n'
    '三、管理员审核与核验流程：\n'
    '  7. 前端 AdminDashboard → adoptAPI.audit(id, data) → PUT /api/adopt/applications/{id}/audit/\n'
    '     → AdminAdoptApplicationViewSet.update() → [IsAdminRole]校验\n'
    '       → AdoptApplicationAuditSerializer(data).is_valid()\n'
    '         → validate(): 拒绝时audit_opinion必填\n'
    '       → serializer.save() → update() 重写:\n'
    '         → instance.auditor = request.user\n'
    '         → instance.audited_at = timezone.now()\n'
    '         → 若approved → instance.pet.adoption_status = "pending"\n'
    '         → 若rejected且无其他活跃申请 → instance.pet.adoption_status = "available"\n'
    '       → write_operation_log(user, "adopt", "audit", ...) 记录到 system.OperationLog\n\n'
    '  8. adoptAPI.offlineVerify(id, data) → PUT /api/adopt/offline-verify/{id}/\n'
    '     → AdminOfflineVerifyViewSet.update() → [IsAdminRole]校验\n'
    '       → 若verify_status="failed"且verify_note为空 → 400\n'
    '       → verify.verifier = request.user, verify.verified_at = timezone.now()\n'
    '       → 若passed → application.online_status="approved" + pet.adoption_status="adopted"\n'
    '       → 若failed → application.online_status="rejected" + pet.adoption_status="available"\n'
)

doc.add_heading('1.2.2  前端组件调用链', level=3)

doc.add_paragraph(
    '一、用户浏览与申请流程：\n'
    '  App.js (路由) → /pets → PetList 组件\n'
    '    → petsAPI.getAll({adoption_status:"available", ...筛选参数})\n'
    '    → 点击卡片 → navigate(/pets/:id) → PetDetail 组件\n'
    '      → petsAPI.getById(id) + petsAPI.getAll()（前后切换导航用）\n'
    '      → 点击"带我回家"按钮 → requireAuth() 检查登录\n'
    '        → 若未登录 → navigate(/login)\n'
    '        → 若已登录 → navigate(/adopt/:petId) → AdoptApplication 组件\n'
    '          步骤1: 填写16字段表单 → validateStep1() → adoptAPI.create({pet_id, message})\n'
    '            → adoptAPI.submitQuestionnaire(appId, answers) → setStep(2)\n'
    '          步骤2: 上传身份证+居住证明 → uploadAPI.upload(file, "adopt")\n'
    '            → adoptAPI.addAttachment(appId, {file_type, file_url, file_size})\n'
    '            → 全部上传完成 → setStep(3) 显示成功页\n\n'
    '二、用户查看申请状态流程：\n'
    '  App.js → /my-applications → MyApplications 组件\n'
    '    → adoptAPI.getMy() → 获取所有申请\n'
    '    → getDisplayStatus(app) 综合 online_status + verify_status 推断显示状态\n'
    '    → 仅"审核拒绝"或"核验失败"可点击"查看详情"\n'
    '    → /my-applications/:id → ApplicationDetail 组件\n'
    '      → adoptAPI.getById(id) → getDetailType(app) → 显示驳回/核验失败原因\n\n'
    '三、管理员管理流程：\n'
    '  App.js → /add-pet → AddPet 组件（ProtectedRoute + admin角色检查）\n'
    '    → authAPI.getProfile() 验证 role==="admin"\n'
    '    → rescueAPI.getAll() 加载救助案例列表（用于关联）\n'
    '    → AI功能: aiAPI.breedDetect() / aiAPI.adoptCopy()\n'
    '    → petsAPI.create(payload) → 创建宠物档案\n\n'
    '  App.js → /admin → AdminDashboard 组件（ProtectedRoute + AdminRoute）\n'
    '    领养审核Tab:\n'
    '      → adoptAPI.getAll() → 表格展示所有申请\n'
    '      → adoptAPI.audit(id, {online_status, audit_opinion}) → 审核操作\n'
    '    宠物管理Tab:\n'
    '      → petsAPI.getAll() → 卡片列表\n'
    '      → petsAPI.update(id, {is_public}) / petsAPI.delete(id)\n'
)

doc.add_page_break()

# ---- 1.3 源代码 ----
doc.add_heading('1.3  源代码', level=2)

doc.add_paragraph('以下为宠物领养模块的完整源代码，按文件路径排列，关键部分已添加中文注释。')

# 后端代码
doc.add_heading('1.3.1  后端源代码', level=3)

# models.py
add_code_block(doc, 'backend/pets/models.py', read_file('backend/pets/models.py'))

# serializers.py
add_code_block(doc, 'backend/pets/serializers.py', read_file('backend/pets/serializers.py'))

# views.py
add_code_block(doc, 'backend/pets/views.py', read_file('backend/pets/views.py'))

# urls.py
add_code_block(doc, 'backend/pets/urls.py', read_file('backend/pets/urls.py'))

# adopt_urls.py
add_code_block(doc, 'backend/pets/adopt_urls.py', read_file('backend/pets/adopt_urls.py'))

# common/permissions.py
add_code_block(doc, 'backend/common/permissions.py', read_file('backend/common/permissions.py'))

# 前端代码
doc.add_heading('1.3.2  前端源代码', level=3)

# API文件
add_code_block(doc, 'frontend/src/api/adopt.js', read_file('frontend/src/api/adopt.js'))
add_code_block(doc, 'frontend/src/api/pets.js', read_file('frontend/src/api/pets.js'))
add_code_block(doc, 'frontend/src/constants/site.js', read_file('frontend/src/constants/site.js'))

# 页面组件
add_code_block(doc, 'frontend/src/pages/PetList.js', read_file('frontend/src/pages/PetList.js'))
add_code_block(doc, 'frontend/src/pages/PetDetail.js', read_file('frontend/src/pages/PetDetail.js'))
add_code_block(doc, 'frontend/src/pages/AdoptApplication.js', read_file('frontend/src/pages/AdoptApplication.js'))
add_code_block(doc, 'frontend/src/pages/MyApplications.js', read_file('frontend/src/pages/MyApplications.js'))
add_code_block(doc, 'frontend/src/pages/ApplicationDetail.js', read_file('frontend/src/pages/ApplicationDetail.js'))
add_code_block(doc, 'frontend/src/pages/AddPet.js', read_file('frontend/src/pages/AddPet.js'))

doc.add_page_break()

# ======================================================================
#              2  救助追踪功能
# ======================================================================
doc.add_heading('2  救助追踪功能', level=1)

doc.add_paragraph(
    '救助追踪模块负责流浪动物发现上报、救助人响应、状态流转追踪和阶段记录管理。'
    '该模块实现了完整的状态机流转机制，确保救助过程有序推进。'
    '由后端 Django 应用 rescue 和前端 React 页面共同组成。'
)

# ---- 2.1 实现类 ----
doc.add_heading('2.1  实现类（类保护的属性与方法）', level=2)

doc.add_heading('2.1.1  后端模型类（Models）', level=3)

add_table(doc,
    ['类名', '继承自', '主要属性/字段', '访问修饰符', '文件路径'],
    [
        ['RescueCase',
         'models.Model',
         'STATUS_CHOICES (类常量, 6种状态)\n'
         'SIZE_CHOICES (类常量)\n'
         'HEALTH_CHOICES (类常量)\n'
         'rescue_no: CharField(32, unique) — 救助编号\n'
         'reporter: ForeignKey→User (RESTRICT)\n'
         'helpers: ManyToManyField→User (blank=True)\n'
         'help_date: DateTimeField\n'
         'nickname: CharField(50)\n'
         'contact: CharField(100)\n'
         'discover_latitude: DecimalField(9,6)\n'
         'discover_longitude: DecimalField(9,6)\n'
         'discover_address: CharField(255)\n'
         'size_category: CharField(20)\n'
         'health_status: CharField(20)\n'
         'is_injured: BooleanField\n'
         'afraid_of_people: BooleanField\n'
         'appearance: TextField\n'
         'health_note: TextField\n'
         'photo_urls: JSONField(default=list)\n'
         'current_status: CharField(20)\n'
         'Meta: db_table="rescue_case"',
         'public',
         'backend/rescue/models.py'],

        ['RescueStatusLog',
         'models.Model',
         'rescue_case: ForeignKey→RescueCase (CASCADE)\n'
         'from_status: CharField(20) — 变更前状态\n'
         'to_status: CharField(20) — 变更后状态\n'
         'operator: ForeignKey→User (RESTRICT)\n'
         'remark: TextField — 操作备注（必填）\n'
         'created_at: DateTimeField\n'
         'Meta: db_table="rescue_status_log"',
         'public',
         'backend/rescue/models.py'],

        ['RescueStageRecord',
         'models.Model',
         'rescue_case: ForeignKey→RescueCase (CASCADE)\n'
         'content: TextField — 阶段详细记录\n'
         'operator: ForeignKey→User (RESTRICT)\n'
         'created_at: DateTimeField\n'
         'Meta: db_table="rescue_stage_record"',
         'public',
         'backend/rescue/models.py'],
    ]
)

doc.add_heading('2.1.2  后端序列化器类（Serializers）', level=3)

add_table(doc,
    ['类名/函数名', '继承自', '主要属性/方法', '访问修饰符', '文件路径'],
    [
        ['_round_coordinate(value)',
         '模块级函数',
         '将坐标值四舍五入到6位小数\n返回 Decimal 类型',
         '模块私有（_前缀）',
         'backend/rescue/serializers.py'],

        ['CoordinateField',
         'serializers.DecimalField',
         '__init__: max_digits=9, decimal_places=6\n'
         'to_internal_value(self, data)\n'
         '  → 调用 _round_coordinate() 舍入后转换',
         'public',
         'backend/rescue/serializers.py'],

        ['RescueStatusLogSerializer',
         'serializers.ModelSerializer',
         'operator: UserSerializer(read_only)\n'
         'Meta: model=RescueStatusLog, fields="__all__"',
         'public',
         'backend/rescue/serializers.py'],

        ['RescueCaseSerializer',
         'serializers.ModelSerializer',
         'reporter: UserSerializer(read_only)\n'
         'helpers: UserSerializer(many=True, read_only)\n'
         'status_logs: RescueStatusLogSerializer(many, read_only)\n'
         'discover_latitude/longitude: CoordinateField()\n'
         'validate_discover_address(value) — 非空校验\n'
         'validate_nickname(value) — 非空校验\n'
         'validate_contact(value) — 非空且≥5字符\n'
         'read_only_fields: [rescue_no,reporter,helpers,help_date,created_at,updated_at]',
         'public',
         'backend/rescue/serializers.py'],

        ['RescueStageRecordSerializer',
         'serializers.ModelSerializer',
         'operator: UserSerializer(read_only)\n'
         'read_only_fields: [rescue_case,operator,created_at]',
         'public',
         'backend/rescue/serializers.py'],

        ['generate_rescue_no()',
         '模块级函数',
         '生成格式为 RES{YYYYMMDD}{NNN} 的唯一编号\n'
         '查询当日最大编号后递增流水号',
         'public（模块导出）',
         'backend/rescue/serializers.py'],
    ]
)

doc.add_heading('2.1.3  后端视图类（Views）', level=3)

add_table(doc,
    ['类名', '继承自', '主要方法', '权限控制', '文件路径'],
    [
        ['RescueCasePagination',
         'PageNumberPagination',
         'page_size = 10\n'
         'page_size_query_param = "page_size"\n'
         'max_page_size = 50',
         '— (分页配置类)',
         'backend/rescue/views.py'],

        ['RescueCaseViewSet',
         'viewsets.ModelViewSet',
         'STATUS_FLOW = {\n'
         '  "in_medical":"recovering",\n'
         '  "recovering":"awaiting_adoption",\n'
         '  "awaiting_adoption":"rescued"\n'
         '}\n'
         'get_queryset(self) — 支持my/helped/status/\n'
         '  rescue_no/exclude_status筛选\n'
         'get_permissions(self)\n'
         '  list/retrieve: AllowAny\n'
         '  create/help/update_status/stage_records: IsAuthenticated\n'
         '  其他: IsAdminRole\n'
         'perform_create(self, serializer)\n'
         '  — 设置reporter+rescue_no+初始状态日志\n'
         'help(self, request, pk) — POST 响应救助\n'
         '  — 自动推进状态 pending_rescue→in_medical\n'
         'update_status(self, request, pk) — PATCH 更新状态\n'
         '  — 按STATUS_FLOW规则校验+备注必填\n'
         'stage_records(self, request, pk) — GET/POST\n'
         '  — 查看/新增阶段记录',
         '分层权限',
         'backend/rescue/views.py'],
    ]
)

doc.add_heading('2.1.4  前端组件类', level=3)

add_table(doc,
    ['组件/模块名', '类型', '主要属性/方法', '路由', '文件路径'],
    [
        ['rescueAPI',
         'API模块（对象）',
         'getAll(params), getById(id), getMyCases(),\n'
         'getHelpedCases(params), create(data),\n'
         'help(id), updateStatus(id,data),\n'
         'getStageRecords(id), addStageRecord(id,data)',
         '—',
         'frontend/src/api/pets.js'],

        ['RescueList',
         'React Function Component',
         'State: cases,loading,error,helpingId\n'
         'localStorage: pet_connect_my_rescues\n'
         'Methods: loadMyRescues(),saveMyRescues(),\n'
         '  buildDescription(),handleHelp(),\n'
         '  stripCoord()',
         '/rescue',
         'frontend/src/pages/RescueList.js'],

        ['RescueReport',
         'React Function Component',
         'State: form(9字段),photos,uploading,\n'
         '  submitting,locating,locationHint,error,\n'
         '  fieldErrors,successNo\n'
         'Methods: handleUseCurrentLocation(),\n'
         '  reverseGeocode(lat,lng) — 高德逆地理编码,\n'
         '  handlePhotoSelect(),validate(),\n'
         '  handleSubmit(),handleReset()',
         '/rescue/report',
         'frontend/src/pages/RescueReport.js'],

        ['RescueSearch',
         'React Function Component',
         'State: keyword,results,searching,searched\n'
         'Methods: handleSearch(),handleKeyDown()',
         '/rescue/search',
         'frontend/src/pages/RescueSearch.js'],

        ['RescueDetail',
         'React Function Component',
         'State: caseData,stageRecords,loading,\n'
         '  error,activeStatus\n'
         'Methods: fetchData(),构建timelineByStatus\n'
         '  (合并status_logs+stage_records),\n'
         '  formatDateTime(),STATUS_META元信息',
         '/rescue/:id',
         'frontend/src/pages/RescueDetail.js'],

        ['MyRescues',
         'React Function Component',
         'State: records,loading,error,statusFilter,\n'
         '  page,totalPages,totalCount\n'
         'Methods: fetchHelped(),handleFilterChange(),\n'
         '  handleReset(),formatDate(),showUpdateBtn()',
         '/my-rescues',
         'frontend/src/pages/MyRescues.js'],

        ['UpdateRescueStatus',
         'React Function Component',
         'State: caseData,newStatus,remark,loading,\n'
         '  submitting,error,success,validationError\n'
         'STATUS_NEXT: in_medical→recovering→\n'
         '  awaiting_adoption→rescued\n'
         'Methods: handleSwitch(),handleSubmit()',
         '/my-rescues/:id/update-status',
         'frontend/src/pages/UpdateRescueStatus.js'],

        ['StageRecord',
         'React Function Component',
         'State: caseData,records,content,loading,\n'
         '  submitting,error,success,validationError\n'
         'Methods: fetchData(),handleSubmit()',
         '/my-rescues/:id/stage-record',
         'frontend/src/pages/StageRecord.js'],
    ]
)

doc.add_page_break()

# ---- 2.2 调用关系 ----
doc.add_heading('2.2  各个类及方法的调用关系', level=2)

doc.add_heading('2.2.1  后端调用链', level=3)

doc.add_paragraph(
    '一、救助案例创建与上报流程：\n'
    '  1. 前端 RescueReport 组件 → rescueAPI.create(data) → POST /api/rescue/cases/\n'
    '     → urls.py (DefaultRouter) → RescueCaseViewSet.create()\n'
    '     → get_permissions() → [IsAuthenticated]\n'
    '     → RescueCaseSerializer(data).is_valid()\n'
    '       → validate_nickname(value) 非空校验\n'
    '       → validate_contact(value) 非空且≥5字符\n'
    '       → validate_discover_address(value) 非空校验\n'
    '       → CoordinateField.to_internal_value() 坐标舍入\n'
    '     → perform_create(serializer)\n'
    '       → generate_rescue_no() 生成 RES{YYYYMMDD}{NNN}\n'
    '       → serializer.save(reporter=request.user, rescue_no=...)\n'
    '       → RescueStatusLog.objects.create(\n'
    '            from_status=None, to_status="pending_rescue",\n'
    '            remark="被发现上报")\n'
    '     → Response 201 (含 rescue_no)\n\n'
    '二、救助案例查询流程：\n'
    '  2. 前端 RescueList → rescueAPI.getMyCases() → GET /api/rescue/cases/?my=true&exclude_status=awaiting_adoption\n'
    '     → RescueCaseViewSet.list() → get_queryset()\n'
    '       → filter(reporter=request.user)\n'
    '       → exclude(current_status__in=["awaiting_adoption"])\n'
    '       → .distinct()\n'
    '     → RescueCasePagination 分页\n'
    '     → RescueCaseSerializer 序列化（含 status_logs 嵌套）\n\n'
    '  3. 前端 MyRescues → rescueAPI.getHelpedCases(params) → GET /api/rescue/cases/?helped=true&...\n'
    '     → get_queryset() → filter(helpers=request.user) + 分页\n\n'
    '  4. 前端 RescueSearch → rescueAPI.getAll({rescue_no}) → GET /api/rescue/cases/?rescue_no=xxx\n'
    '     → get_queryset() → filter(rescue_no__iexact=rescue_no)\n\n'
    '三、救助响应与状态流转流程：\n'
    '  5. 前端 RescueList → rescueAPI.help(id) → POST /api/rescue/cases/{id}/help/\n'
    '     → RescueCaseViewSet.help() → [IsAuthenticated]\n'
    '       → 检查重复响应 case.helpers.filter(id=user.id).exists()\n'
    '       → case.helpers.add(user)\n'
    '       → 首次响应: case.help_date = datetime.now()\n'
    '       → 状态流转: pending_rescue → in_medical\n'
    '       → case.save(update_fields=[...])\n'
    '       → RescueStatusLog.objects.create(from_status, to_status, operator, remark)\n'
    '       → write_operation_log(user, "rescue", "help", ...)\n'
    '     → Response 200 (序列化后的案例数据)\n\n'
    '  6. 前端 UpdateRescueStatus → rescueAPI.updateStatus(id, {current_status, remark})\n'
    '     → PATCH /api/rescue/cases/{id}/status/\n'
    '     → RescueCaseViewSet.update_status() → [IsAuthenticated]\n'
    '       → 权限检查: is_helper OR is_admin，否则403\n'
    '       → 校验: new_status ∈ STATUS_CHOICES\n'
    '       → 校验: remark 非空\n'
    '       → 校验: 当前状态在 STATUS_FLOW 中（非终态）\n'
    '       → 校验: new_status == STATUS_FLOW[old_status]（只能顺序推进）\n'
    '       → case.current_status = new_status\n'
    '       → RescueStatusLog.objects.create(...)\n'
    '       → write_operation_log(...)\n'
    '     → Response 200\n\n'
    '四、阶段记录流程：\n'
    '  7. 前端 StageRecord → rescueAPI.addStageRecord(id, {content})\n'
    '     → POST /api/rescue/cases/{id}/stage-records/\n'
    '     → RescueCaseViewSet.stage_records() (POST分支)\n'
    '       → 校验 content 非空\n'
    '       → 权限检查: is_helper OR is_admin\n'
    '       → RescueStageRecord.objects.create(rescue_case, content, operator)\n'
    '     → Response 201\n\n'
    '  8. 前端 RescueDetail → rescueAPI.getStageRecords(id)\n'
    '     → GET /api/rescue/cases/{id}/stage-records/\n'
    '     → RescueCaseViewSet.stage_records() (GET分支)\n'
    '       → case.stage_records.select_related("operator").all()\n'
    '       → RescueStageRecordSerializer 序列化\n'
    '     → Response 200\n'
)

doc.add_heading('2.2.2  前端组件调用链', level=3)

doc.add_paragraph(
    '一、浏览与响应救助流程：\n'
    '  App.js → /rescue → RescueList 组件 (ProtectedRoute)\n'
    '    → rescueAPI.getMyCases() → 获取待救助列表（排除"待领养"状态）\n'
    '    → localStorage: loadMyRescues() 加载已关注列表\n'
    '    → 过滤已关注的案例 → 展示剩余待救助列表\n'
    '    → 点击"救助"按钮 → handleHelp(item)\n'
    '      → rescueAPI.help(item.id) → 响应救助\n'
    '      → saveMyRescues() 保存到 localStorage\n'
    '      → 300ms后从列表中移除\n'
    '  顶部按钮:\n'
    '    → "我的救助" → navigate(/my-rescues) → MyRescues 组件\n'
    '    → "上报" → navigate(/rescue/report) → RescueReport 组件\n'
    '    → "查询" → navigate(/rescue/search) → RescueSearch 组件\n\n'
    '二、上报救助流程：\n'
    '  RescueReport 组件:\n'
    '    → handleUseCurrentLocation()\n'
    '      → navigator.geolocation.getCurrentPosition() → 浏览器定位\n'
    '      → reverseGeocode(lat, lng) → 高德地图API逆地理编码\n'
    '      → setForm(discover_latitude, discover_longitude, discover_address)\n'
    '    → handlePhotoSelect()\n'
    '      → 文件类型/大小校验（jpg/png/pdf, ≤5MB）\n'
    '      → uploadAPI.upload(file, "rescue") → 上传照片\n'
    '    → handleSubmit()\n'
    '      → validate() 9字段校验\n'
    '      → rescueAPI.create(payload) → 创建救助案例\n'
    '      → setSuccessNo(res.data.rescue_no) → 显示成功页\n\n'
    '三、救助追踪流程：\n'
    '  MyRescues 组件:\n'
    '    → rescueAPI.getHelpedCases({page, status}) → 分页获取我参与的案例\n'
    '    → 按状态筛选（下拉框选择）\n'
    '    → 每条记录显示: 照片/编号/状态/日期\n'
    '    → "更新状态"按钮 → navigate(/my-rescues/:id/update-status)\n'
    '    → "填写记录"按钮 → navigate(/my-rescues/:id/stage-record)\n\n'
    '  UpdateRescueStatus 组件:\n'
    '    → rescueAPI.getById(id) → 加载案例数据\n'
    '    → handleSwitch() → 计算 STATUS_NEXT[currentStatus] 下一状态\n'
    '    → handleSubmit() → rescueAPI.updateStatus(id, {current_status, remark})\n'
    '    → 只能按 in_medical→recovering→awaiting_adoption→rescued 顺序推进\n\n'
    '  StageRecord 组件:\n'
    '    → Promise.all([rescueAPI.getById(id), rescueAPI.getStageRecords(id)])\n'
    '    → handleSubmit() → rescueAPI.addStageRecord(id, {content})\n'
    '    → 新增记录插入列表顶部\n'
    '    → 历史记录列表按时间倒序展示\n\n'
    '四、救助详情与时间线：\n'
    '  RescueDetail 组件:\n'
    '    → Promise.all([rescueAPI.getById(id), rescueAPI.getStageRecords(id)])\n'
    '    → 从 status_logs 构建 reachedStatuses（已达状态列表）\n'
    '    → 构建 timelineByStatus:\n'
    '      → 合并 status_logs（状态变更） + stage_records（阶段记录）\n'
    '      → 每条阶段记录按时间归属到对应状态区间\n'
    '      → 按时间排序\n'
    '    → 展示状态流转路径（可点击切换查看不同状态的记录）\n'
    '    → 查看详情 → navigate(/rescue/:id)（从列表或搜索结果进入）\n'
)

doc.add_page_break()

# ---- 2.3 源代码 ----
doc.add_heading('2.3  源代码', level=2)

doc.add_paragraph('以下为救助追踪模块的完整源代码，按文件路径排列，关键部分已添加中文注释。')

doc.add_heading('2.3.1  后端源代码', level=3)

# models.py
add_code_block(doc, 'backend/rescue/models.py', read_file('backend/rescue/models.py'))

# serializers.py
add_code_block(doc, 'backend/rescue/serializers.py', read_file('backend/rescue/serializers.py'))

# views.py
add_code_block(doc, 'backend/rescue/views.py', read_file('backend/rescue/views.py'))

# urls.py
add_code_block(doc, 'backend/rescue/urls.py', read_file('backend/rescue/urls.py'))

doc.add_heading('2.3.2  前端源代码', level=3)

# rescueAPI 在 pets.js 中，已经贴过，这里标注引用
doc.add_paragraph('注：rescueAPI 和 uploadAPI 均定义在 frontend/src/api/pets.js 中，该文件已在 1.3.2 节完整列出。')
doc.add_paragraph('constants/site.js 中的 RESCUE_STATUS / SIZE_CATEGORY / HEALTH_STATUS 常量也已在 1.3.2 节列出。')

add_code_block(doc, 'frontend/src/pages/RescueList.js', read_file('frontend/src/pages/RescueList.js'))
add_code_block(doc, 'frontend/src/pages/RescueReport.js', read_file('frontend/src/pages/RescueReport.js'))
add_code_block(doc, 'frontend/src/pages/RescueSearch.js', read_file('frontend/src/pages/RescueSearch.js'))
add_code_block(doc, 'frontend/src/pages/RescueDetail.js', read_file('frontend/src/pages/RescueDetail.js'))
add_code_block(doc, 'frontend/src/pages/MyRescues.js', read_file('frontend/src/pages/MyRescues.js'))
add_code_block(doc, 'frontend/src/pages/UpdateRescueStatus.js', read_file('frontend/src/pages/UpdateRescueStatus.js'))
add_code_block(doc, 'frontend/src/pages/StageRecord.js', read_file('frontend/src/pages/StageRecord.js'))

# ===== 保存文档 =====
output_path = os.path.join(BASE_DIR, 'docs', '宠物领养与救助追踪模块编码文档.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
